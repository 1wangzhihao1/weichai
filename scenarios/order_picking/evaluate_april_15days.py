import argparse
import json
import os
import sys
from collections import defaultdict
from datetime import date
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd

CURRENT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = CURRENT_DIR.parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from backend.replay_history import (
    DEFAULT_GLOBAL_BREAK_THRESHOLD_SECONDS,
    DEFAULT_PICK_GAP_SECONDS,
    calculate_observed_metrics,
    get_db_part_times,
    load_historical_orders,
    replay_historical_data,
)
from scenarios.order_picking.compare import (
    load_latest_model,
    order_date,
    run_simulation,
    set_daily_orders,
)
from scenarios.order_picking.config import Config
from scenarios.order_picking.data_paths import OUTPUT_DIR, historical_picking_excel
from scenarios.order_picking.rl_environment import PickingEnv


plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "Arial Unicode MS", "sans-serif"]
plt.rcParams["axes.unicode_minus"] = False

ORDER_ID_COLUMN = "拣选列表"
OUTPUT_ROOT = OUTPUT_DIR / "april_15day_validation"


def parse_args():
    parser = argparse.ArgumentParser(
        description="Validate AI scheduling on 15 April dates against rule baselines and historical replay."
    )
    parser.add_argument("--date-count", type=int, default=15, help="Number of April dates to evaluate.")
    parser.add_argument("--month-prefix", default="2026-04", help="Month prefix, for example 2026-04.")
    parser.add_argument("--selection", choices=["top_orders", "first"], default="top_orders")
    parser.add_argument("--min-stations", type=int, default=1)
    parser.add_argument("--max-stations", type=int, default=Config.NUM_STATIONS)
    parser.add_argument("--seed", type=int, default=888)
    parser.add_argument("--pick-gap-seconds", type=float, default=DEFAULT_PICK_GAP_SECONDS)
    parser.add_argument("--skip-docx", action="store_true", help="Generate JSON/XLSX/charts only.")
    parser.add_argument(
        "--global-break-threshold-seconds",
        type=float,
        default=DEFAULT_GLOBAL_BREAK_THRESHOLD_SECONDS,
    )
    return parser.parse_args()


def pct_improve(baseline, candidate):
    if baseline is None or candidate is None or baseline <= 0:
        return None
    return round((baseline - candidate) / baseline * 100.0, 3)


def station_hour_improve(baseline_stations, baseline_seconds, ai_stations, ai_seconds):
    if (
        baseline_stations is None
        or baseline_seconds is None
        or ai_stations is None
        or ai_seconds is None
        or baseline_stations <= 0
        or baseline_seconds <= 0
    ):
        return None
    baseline_station_hours = float(baseline_stations) * float(baseline_seconds)
    ai_station_hours = float(ai_stations) * float(ai_seconds)
    return round((1.0 - ai_station_hours / baseline_station_hours) * 100.0, 3)


def seconds_to_hours(value):
    return round(float(value) / 3600.0, 3) if value is not None else None


def load_history_frame():
    path = historical_picking_excel()
    if not path.exists():
        raise FileNotFoundError(f"Historical picking Excel not found: {path}")
    print(f"Reading historical Excel: {path}")
    df = pd.read_excel(path, sheet_name=0)
    df.columns = df.columns.astype(str).str.strip().str.replace("\n", "").str.replace("\r", "")
    df["开始时间"] = pd.to_datetime(df["开始时间"], errors="coerce")
    df["结束时间"] = pd.to_datetime(df["结束时间"], errors="coerce")
    df["测试日期"] = df["开始时间"].dt.strftime("%Y-%m-%d")
    return df


def station_id(value):
    if pd.isna(value):
        return -1
    digits = "".join(ch for ch in str(value) if ch.isdigit())
    return int(digits) if digits else -1


def valid_history_by_date(history_df, month_prefix):
    df = history_df[history_df["测试日期"].astype(str).str.startswith(month_prefix)].copy()
    if "状态" in df.columns:
        df = df[df["状态"].astype(str).str.strip().isin(["完成", "确定"])]
    df = df.dropna(subset=["开始时间", "结束时间", "SKU", ORDER_ID_COLUMN, "拣选员 ID", "已拣选数量"])
    df["已拣选数量"] = pd.to_numeric(df["已拣选数量"], errors="coerce")
    df = df[df["已拣选数量"] > 0]
    return {day: group.copy() for day, group in df.groupby("测试日期")}


def load_orders_by_date():
    env = PickingEnv(dataset_type="all", enable_inventory=False)
    orders = list(env.unwrapped.real_world_orders)
    orders_by_date = defaultdict(list)
    for order in orders:
        orders_by_date[order_date(order)].append(order)
    return env, orders_by_date


def select_dates(orders_by_date, history_by_date, args):
    candidates = []
    for day in sorted(set(orders_by_date) & set(history_by_date)):
        if not day.startswith(args.month_prefix):
            continue
        candidates.append(
            {
                "date": day,
                "order_count": len(orders_by_date[day]),
                "history_rows": len(history_by_date[day]),
            }
        )
    if len(candidates) < args.date_count:
        raise RuntimeError(f"Only {len(candidates)} matched dates found, less than {args.date_count}.")

    if args.selection == "top_orders":
        selected = sorted(candidates, key=lambda row: (-row["order_count"], row["date"]))[: args.date_count]
        return sorted(row["date"] for row in selected)
    return [row["date"] for row in candidates[: args.date_count]]


def best_record(strategy_rows, strategy_key):
    valid = [row for row in strategy_rows if row[strategy_key] <= Config.DEADLINE_SECONDS]
    if valid:
        return min(valid, key=lambda row: row["station_limit"])
    return min(strategy_rows, key=lambda row: row[strategy_key])


def run_strategy_scans(env, model, orders, station_limits, seed):
    rows = []
    set_daily_orders(env, orders)
    for limit in station_limits:
        rr = run_simulation(env, "round_robin", limit, seed=seed + limit)
        random_ms = run_simulation(env, "random", limit, seed=seed + limit + 10000)
        ai_ms = run_simulation(env, "ai", limit, model=model, seed=seed + limit)
        rows.append(
            {
                "station_limit": int(limit),
                "ai_makespan_sec": round(float(ai_ms), 3),
                "round_robin_makespan_sec": round(float(rr), 3),
                "random_makespan_sec": round(float(random_ms), 3),
            }
        )
    return rows


def historical_station_summary(day_history):
    df = day_history.copy()
    df["station_id"] = df["拣选员 ID"].apply(station_id)
    df = df[df["station_id"] > 0]
    historical_station_count = int(df["station_id"].nunique())
    order_count = int(df[ORDER_ID_COLUMN].nunique())
    box_count = int(len(df))
    return historical_station_count, order_count, box_count


def make_chart_daily(results, key_a, key_b, label_a, label_b, title, out_path):
    dates = [row["date"] for row in results]
    x = np.arange(len(dates))
    width = 0.38
    a = [row[key_a] / 3600.0 for row in results]
    b = [row[key_b] / 3600.0 for row in results]

    plt.figure(figsize=(13, 5.8))
    plt.bar(x - width / 2, a, width, label=label_a, color="#2E74B5")
    plt.bar(x + width / 2, b, width, label=label_b, color="#A5A5A5")
    plt.xticks(x, dates, rotation=45, ha="right")
    plt.ylabel("完工时间/工作时间（小时）")
    plt.title(title)
    plt.grid(axis="y", alpha=0.25)
    plt.legend()
    plt.tight_layout()
    plt.savefig(out_path, dpi=220)
    plt.close()
    return str(out_path)


def make_chart_strategy(results, out_path):
    dates = [row["date"] for row in results]
    x = np.arange(len(dates))
    plt.figure(figsize=(13, 5.8))
    plt.plot(x, [row["ai_best_makespan_sec"] / 3600.0 for row in results], marker="o", label="AI", linewidth=2.5)
    plt.plot(
        x,
        [row["round_robin_best_makespan_sec"] / 3600.0 for row in results],
        marker="s",
        label="轮询",
        linewidth=2,
    )
    plt.plot(
        x,
        [row["random_best_makespan_sec"] / 3600.0 for row in results],
        marker="x",
        label="随机",
        linewidth=2,
    )
    plt.xticks(x, dates, rotation=45, ha="right")
    plt.ylabel("完工时间（小时）")
    plt.title("AI、轮询、随机策略每日最优完工时间对比")
    plt.grid(axis="y", alpha=0.25)
    plt.legend()
    plt.tight_layout()
    plt.savefig(out_path, dpi=220)
    plt.close()
    return str(out_path)


def set_cell_shading(cell, fill):
    ensure_docx_imports()
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_widths(table, widths):
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)


def style_table(table):
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run.font.size = Pt(8.5)
            if row_idx == 0:
                set_cell_shading(cell, "E8EEF5")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = str(header)
    for row_values in rows:
        row = table.add_row()
        for cell, value in zip(row.cells, row_values):
            cell.text = "" if value is None else str(value)
    style_table(table)
    if widths:
        set_table_widths(table, widths)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    for style_name, size, color in [
        ("Heading 1", 15, RGBColor(46, 116, 181)),
        ("Heading 2", 12, RGBColor(31, 77, 120)),
        ("Heading 3", 11, RGBColor(31, 77, 120)),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def ensure_docx_imports():
    global Document, WD_ALIGN_PARAGRAPH, OxmlElement, qn, Inches, Pt, RGBColor
    if "Document" in globals():
        return
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor


def add_summary_paragraph(doc, summary):
    p = doc.add_paragraph()
    p.add_run("核心结论：").bold = True
    p.add_run(
        f"本次从四月份抽取 {summary['date_count']} 天进行验证。"
        f"AI 相比轮询平均提升 {summary['avg_ai_vs_round_robin_improvement_pct']:.2f}%，"
        f"相比随机平均提升 {summary['avg_ai_vs_random_improvement_pct']:.2f}%。"
        f"在同历史站台数口径下，AI 相比历史仿真平均提升 "
        f"{summary['avg_ai_vs_history_replay_same_station_improvement_pct']:.2f}%。"
    )


def generate_docx(results, summary, charts, output_path):
    ensure_docx_imports()
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("四月份十五天订单分拣调度验证报告")
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("AI 调度、随机/轮询策略、历史真实时间与历史仿真时间对比").italic = True

    doc.add_heading("1. 验证口径", level=1)
    for text in [
        "数据范围：DMS拣选20260201-0429.XLSX 中四月份抽取的 15 天有效拣选数据。",
        "库存口径：本次所有策略和历史回放均不考虑库存因素。",
        "策略对比：AI、轮询、随机均使用同一批订单和同一套 SKU 标准工时。",
        "历史真实时间：使用 Excel 中真实开始/结束时间，并扣除超过 30 分钟的全局无作业间隔。",
        "历史仿真时间：使用历史真实站台分配和订单顺序，但处理时间采用数据库 PartMaster 标准工时。",
    ]:
        doc.add_paragraph(text, style=None)

    add_summary_paragraph(doc, summary)

    doc.add_heading("2. 样本日期概况", level=1)
    add_table(
        doc,
        ["日期", "算法订单数", "历史订单数", "历史拣选记录", "历史站台数"],
        [
            [
                r["date"],
                r["algorithm_order_count"],
                r["history_order_count"],
                r["history_box_count"],
                r["historical_station_count"],
            ]
            for r in results
        ],
        [1.0, 1.1, 1.1, 1.1, 1.0],
    )

    doc.add_heading("3. AI、轮询、随机策略对比", level=1)
    doc.add_picture(charts["strategy"], width=Inches(6.7))
    add_table(
        doc,
        ["日期", "AI站台", "AI(h)", "轮询站台", "轮询(h)", "随机站台", "随机(h)", "AI较轮询", "AI较随机"],
        [
            [
                r["date"],
                r["ai_best_stations"],
                seconds_to_hours(r["ai_best_makespan_sec"]),
                r["round_robin_best_stations"],
                seconds_to_hours(r["round_robin_best_makespan_sec"]),
                r["random_best_stations"],
                seconds_to_hours(r["random_best_makespan_sec"]),
                f"{r['ai_vs_round_robin_best_improvement_pct']:.2f}%",
                f"{r['ai_vs_random_best_improvement_pct']:.2f}%",
            ]
            for r in results
        ],
        [0.9, 0.65, 0.75, 0.75, 0.75, 0.75, 0.75, 0.85, 0.85],
    )

    doc.add_heading("4. AI 与历史真实工作时间对比", level=1)
    doc.add_picture(charts["actual"], width=Inches(6.7))
    add_table(
        doc,
        ["日期", "真实总跨度(h)", "休息/停工(h)", "真实净工作(h)", "AI同站台(h)", "AI较真实净时间"],
        [
            [
                r["date"],
                seconds_to_hours(r["actual_span_sec"]),
                seconds_to_hours(r["global_break_total_sec"]),
                seconds_to_hours(r["actual_net_work_sec"]),
                seconds_to_hours(r["ai_same_history_station_makespan_sec"]),
                f"{r['ai_vs_actual_net_improvement_pct']:.2f}%",
            ]
            for r in results
        ],
        [0.95, 1.05, 1.05, 1.1, 1.1, 1.1],
    )

    doc.add_heading("5. AI 与历史仿真时间对比", level=1)
    doc.add_picture(charts["replay"], width=Inches(6.7))
    add_table(
        doc,
        ["日期", "历史仿真(h)", "AI同站台(h)", "AI最优(h)", "AI同站台提升", "AI最优提升"],
        [
            [
                r["date"],
                seconds_to_hours(r["history_replay_makespan_sec"]),
                seconds_to_hours(r["ai_same_history_station_makespan_sec"]),
                seconds_to_hours(r["ai_best_makespan_sec"]),
                f"{r['ai_vs_history_replay_same_station_improvement_pct']:.2f}%",
                f"{r['ai_vs_history_replay_best_improvement_pct']:.2f}%",
            ]
            for r in results
        ],
        [0.9, 1.0, 1.0, 1.0, 1.1, 1.1],
    )

    doc.add_heading("6. 总体平均值", level=1)
    add_table(
        doc,
        ["指标", "平均值"],
        [
            ["AI 最优完工时间(h)", seconds_to_hours(summary["avg_ai_best_makespan_sec"])],
            ["轮询最优完工时间(h)", seconds_to_hours(summary["avg_round_robin_best_makespan_sec"])],
            ["随机最优完工时间(h)", seconds_to_hours(summary["avg_random_best_makespan_sec"])],
            ["历史真实净工作时间(h)", seconds_to_hours(summary["avg_actual_net_work_sec"])],
            ["历史仿真完工时间(h)", seconds_to_hours(summary["avg_history_replay_makespan_sec"])],
            ["AI 相比轮询平均提升", f"{summary['avg_ai_vs_round_robin_improvement_pct']:.2f}%"],
            ["AI 相比随机平均提升", f"{summary['avg_ai_vs_random_improvement_pct']:.2f}%"],
            [
                "AI 同站台相比历史仿真平均提升",
                f"{summary['avg_ai_vs_history_replay_same_station_improvement_pct']:.2f}%",
            ],
        ],
        [3.4, 2.0],
    )

    doc.add_paragraph(
        "说明：历史真实时间包含现场人员节奏、休息、等待与异常等实际因素；历史仿真时间则是在当前标准工时与物理约束下对历史分配方式的复算。因此，AI 与历史仿真的对比更适合衡量分配策略本身的改进。"
    )

    doc.save(output_path)
    return output_path


def summarize(results):
    def avg(key):
        values = [row[key] for row in results if row.get(key) is not None]
        return round(float(np.mean(values)), 3) if values else None

    return {
        "date_count": len(results),
        "avg_ai_best_makespan_sec": avg("ai_best_makespan_sec"),
        "avg_round_robin_best_makespan_sec": avg("round_robin_best_makespan_sec"),
        "avg_random_best_makespan_sec": avg("random_best_makespan_sec"),
        "avg_actual_net_work_sec": avg("actual_net_work_sec"),
        "avg_history_replay_makespan_sec": avg("history_replay_makespan_sec"),
        "avg_ai_vs_round_robin_improvement_pct": avg("ai_vs_round_robin_best_improvement_pct"),
        "avg_ai_vs_random_improvement_pct": avg("ai_vs_random_best_improvement_pct"),
        "avg_ai_vs_actual_net_improvement_pct": avg("ai_vs_actual_net_improvement_pct"),
        "avg_ai_vs_history_replay_same_station_improvement_pct": avg(
            "ai_vs_history_replay_same_station_improvement_pct"
        ),
        "avg_ai_vs_history_replay_best_improvement_pct": avg("ai_vs_history_replay_best_improvement_pct"),
        "avg_ai_vs_round_robin_station_hour_efficiency_pct": avg("ai_vs_round_robin_station_hour_efficiency_pct"),
        "avg_ai_vs_random_station_hour_efficiency_pct": avg("ai_vs_random_station_hour_efficiency_pct"),
        "avg_ai_vs_actual_net_station_hour_efficiency_pct": avg("ai_vs_actual_net_station_hour_efficiency_pct"),
        "avg_ai_vs_history_replay_station_hour_efficiency_pct": avg(
            "ai_vs_history_replay_station_hour_efficiency_pct"
        ),
    }


def save_outputs(results, summary, charts, docx_path):
    OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)
    payload = {
        "generated_on": date.today().isoformat(),
        "deadline_seconds": Config.DEADLINE_SECONDS,
        "results": results,
        "summary": summary,
        "charts": charts,
        "docx": str(docx_path) if docx_path else None,
    }
    json_path = OUTPUT_ROOT / "april_15day_validation_results.json"
    with json_path.open("w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)

    xlsx_path = OUTPUT_ROOT / "april_15day_validation_results.xlsx"
    with pd.ExcelWriter(xlsx_path) as writer:
        pd.DataFrame(results).to_excel(writer, sheet_name="daily_results", index=False)
        pd.DataFrame([summary]).to_excel(writer, sheet_name="summary", index=False)
    return json_path, xlsx_path


def main():
    args = parse_args()
    OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)

    history_df = load_history_frame()
    history_by_date = valid_history_by_date(history_df, args.month_prefix)
    env, orders_by_date = load_orders_by_date()
    selected_dates = select_dates(orders_by_date, history_by_date, args)
    print(f"Selected dates: {selected_dates}")

    model, model_path = load_latest_model(env)
    if model is None:
        raise RuntimeError("AI model is required for validation.")

    part_time_dict = get_db_part_times()
    station_limits = list(range(max(1, args.min_stations), min(args.max_stations, Config.NUM_STATIONS) + 1))
    results = []

    for idx, day in enumerate(selected_dates, start=1):
        print("=" * 80)
        print(f"[{idx}/{len(selected_dates)}] {day}")
        daily_orders = orders_by_date[day]
        day_history = history_by_date[day]
        historical_station_count, history_order_count, history_box_count = historical_station_summary(day_history)
        historical_station_limit = max(1, min(historical_station_count, Config.NUM_STATIONS))

        scans = run_strategy_scans(env, model, daily_orders, station_limits, args.seed + idx * 1000)
        ai_best = best_record(scans, "ai_makespan_sec")
        rr_best = best_record(scans, "round_robin_makespan_sec")
        random_best = best_record(scans, "random_makespan_sec")
        same_station = next(row for row in scans if row["station_limit"] == historical_station_limit)

        observed = calculate_observed_metrics(
            day_history,
            global_break_threshold_seconds=args.global_break_threshold_seconds,
        )
        replay = replay_historical_data(
            target_date=day,
            history_df=history_df,
            part_time_dict=part_time_dict,
            pick_gap_seconds=args.pick_gap_seconds,
            process_time_source="part_master",
            global_break_threshold_seconds=args.global_break_threshold_seconds,
        )
        if replay is None:
            raise RuntimeError(f"History replay failed for {day}")

        row = {
            "date": day,
            "algorithm_order_count": int(len(daily_orders)),
            "history_order_count": history_order_count,
            "history_box_count": history_box_count,
            "historical_station_count": historical_station_count,
            "ai_best_stations": ai_best["station_limit"],
            "ai_best_makespan_sec": ai_best["ai_makespan_sec"],
            "round_robin_best_stations": rr_best["station_limit"],
            "round_robin_best_makespan_sec": rr_best["round_robin_makespan_sec"],
            "random_best_stations": random_best["station_limit"],
            "random_best_makespan_sec": random_best["random_makespan_sec"],
            "ai_same_history_station_makespan_sec": same_station["ai_makespan_sec"],
            "round_robin_same_history_station_makespan_sec": same_station["round_robin_makespan_sec"],
            "random_same_history_station_makespan_sec": same_station["random_makespan_sec"],
            "actual_span_sec": observed.get("actual_span_seconds"),
            "global_break_total_sec": observed.get("global_break_total_seconds"),
            "actual_net_work_sec": observed.get("actual_net_span_excluding_global_breaks_seconds"),
            "history_replay_makespan_sec": replay["historical_replay_makespan_seconds"],
            "history_replay_station_count": replay["historical_station_count"],
            "ai_vs_round_robin_best_improvement_pct": pct_improve(
                rr_best["round_robin_makespan_sec"], ai_best["ai_makespan_sec"]
            ),
            "ai_vs_random_best_improvement_pct": pct_improve(
                random_best["random_makespan_sec"], ai_best["ai_makespan_sec"]
            ),
            "ai_vs_actual_net_improvement_pct": pct_improve(
                observed.get("actual_net_span_excluding_global_breaks_seconds"),
                same_station["ai_makespan_sec"],
            ),
            "ai_vs_history_replay_same_station_improvement_pct": pct_improve(
                replay["historical_replay_makespan_seconds"], same_station["ai_makespan_sec"]
            ),
            "ai_vs_history_replay_best_improvement_pct": pct_improve(
                replay["historical_replay_makespan_seconds"], ai_best["ai_makespan_sec"]
            ),
            "ai_vs_round_robin_station_hour_efficiency_pct": station_hour_improve(
                rr_best["station_limit"],
                rr_best["round_robin_makespan_sec"],
                ai_best["station_limit"],
                ai_best["ai_makespan_sec"],
            ),
            "ai_vs_random_station_hour_efficiency_pct": station_hour_improve(
                random_best["station_limit"],
                random_best["random_makespan_sec"],
                ai_best["station_limit"],
                ai_best["ai_makespan_sec"],
            ),
            "ai_vs_actual_net_station_hour_efficiency_pct": station_hour_improve(
                historical_station_count,
                observed.get("actual_net_span_excluding_global_breaks_seconds"),
                ai_best["station_limit"],
                ai_best["ai_makespan_sec"],
            ),
            "ai_vs_history_replay_station_hour_efficiency_pct": station_hour_improve(
                replay["historical_station_count"],
                replay["historical_replay_makespan_seconds"],
                ai_best["station_limit"],
                ai_best["ai_makespan_sec"],
            ),
            "strategy_scans": scans,
        }
        results.append(row)
        print(
            f"AI best={seconds_to_hours(row['ai_best_makespan_sec'])}h, "
            f"RR best={seconds_to_hours(row['round_robin_best_makespan_sec'])}h, "
            f"Random best={seconds_to_hours(row['random_best_makespan_sec'])}h, "
            f"actual net={seconds_to_hours(row['actual_net_work_sec'])}h, "
            f"history replay={seconds_to_hours(row['history_replay_makespan_sec'])}h"
        )

    summary = summarize(results)
    charts = {
        "strategy": make_chart_strategy(results, OUTPUT_ROOT / "strategy_comparison.png"),
        "actual": make_chart_daily(
            results,
            "ai_same_history_station_makespan_sec",
            "actual_net_work_sec",
            "AI同历史站台数",
            "历史真实净工作时间",
            "AI 与历史真实净工作时间对比",
            OUTPUT_ROOT / "actual_time_comparison.png",
        ),
        "replay": make_chart_daily(
            results,
            "ai_same_history_station_makespan_sec",
            "history_replay_makespan_sec",
            "AI同历史站台数",
            "历史仿真时间",
            "AI 与历史仿真时间对比",
            OUTPUT_ROOT / "history_replay_comparison.png",
        ),
    }
    docx_path = OUTPUT_ROOT / "april_15day_validation_report.docx"
    if not args.skip_docx:
        generate_docx(results, summary, charts, docx_path)
    else:
        docx_path = None
    json_path, xlsx_path = save_outputs(results, summary, charts, docx_path)

    print("=" * 80)
    print(f"Model path : {model_path}")
    print(f"JSON       : {json_path}")
    print(f"Excel      : {xlsx_path}")
    print(f"DOCX       : {docx_path}")
    print("=" * 80)


if __name__ == "__main__":
    main()
