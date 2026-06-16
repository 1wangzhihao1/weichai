import json
import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
TEST_DIR = ROOT / "test"
OUT_DOCX = TEST_DIR / "潍柴订单分拣仿真与调度算法过程留存记录及算法框架图_新增预处理和SKU筛选逻辑版.docx"
DIAGRAM_PATH = TEST_DIR / "algorithm_framework_diagram.png"


def load_json(path):
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


history = load_json(ROOT / "output" / "history_replay_results.json")
comparison = load_json(ROOT / "output" / "comparison_daily_results.json")

history_0411 = next(item for item in history["results"] if item["date"] == "2026-04-11")
comparison_0411 = [item for item in comparison if item["date"] == "2026-04-11"]
comparison_0411_14 = next(item for item in comparison_0411 if item["station_limit"] == 14)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(9)
    run.bold = bold


def style_table(table, header_fill="E8EEF5"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10.5)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10)
    return p


def add_kv_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Cm(4.0)
    table.columns[1].width = Cm(12.0)
    set_cell_text(table.rows[0].cells[0], "项目", True)
    set_cell_text(table.rows[0].cells[1], "内容", True)
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], key, True)
        set_cell_text(cells[1], value)
    style_table(table)
    doc.add_paragraph()
    return table


def add_matrix_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], h, True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    style_table(table)
    doc.add_paragraph()
    return table


def seconds_to_hms(seconds):
    seconds = float(seconds)
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = seconds % 60
    return f"{seconds:.3f} 秒（约 {h}小时{m}分{s:.1f}秒）"


def find_font(size):
    candidates = [
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
    ]
    for candidate in candidates:
        if os.path.exists(candidate):
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def draw_wrapped_text(draw, xy, text, font, fill, max_width, line_gap=6):
    x, y = xy
    lines = []
    current = ""
    for char in text:
        trial = current + char
        bbox = draw.textbbox((0, 0), trial, font=font)
        if bbox[2] - bbox[0] <= max_width or not current:
            current = trial
        else:
            lines.append(current)
            current = char
    if current:
        lines.append(current)
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        draw.text((x + (max_width - (bbox[2] - bbox[0])) / 2, y), line, font=font, fill=fill)
        y += (bbox[3] - bbox[1]) + line_gap


def draw_box(draw, x, y, w, h, title, body, fill, outline):
    draw.rounded_rectangle([x, y, x + w, y + h], radius=18, fill=fill, outline=outline, width=3)
    title_font = find_font(24)
    body_font = find_font(18)
    draw_wrapped_text(draw, (x + 18, y + 18), title, title_font, "#0B2545", w - 36, 5)
    draw_wrapped_text(draw, (x + 18, y + 58), body, body_font, "#243447", w - 36, 5)


def arrow(draw, start, end):
    draw.line([start, end], fill="#51606F", width=4)
    x1, y1 = start
    x2, y2 = end
    if x2 > x1:
        pts = [(x2, y2), (x2 - 14, y2 - 9), (x2 - 14, y2 + 9)]
    elif y2 > y1:
        pts = [(x2, y2), (x2 - 9, y2 - 14), (x2 + 9, y2 - 14)]
    else:
        pts = [(x2, y2), (x2 - 9, y2 + 14), (x2 + 9, y2 + 14)]
    draw.polygon(pts, fill="#51606F")


def create_diagram():
    img = Image.new("RGB", (1800, 1220), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    title_font = find_font(34)
    draw.text((60, 42), "订单分拣强化学习调度算法框架图", font=title_font, fill="#0B2545")

    boxes = {
        "data": (80, 130, 360, 170, "训练数据输入", "历史订单、SKU平均分拣时间、工位基础参数"),
        "clean": (540, 130, 360, 170, "样本构建", "按订单聚合SKU、数量、处理时间和时间顺序"),
        "state": (1000, 130, 360, 170, "状态空间", "站台负载、队列占用、剩余作业量、当前订单特征"),
        "action": (80, 420, 360, 170, "动作空间", "在可用分拣站中选择目标站台"),
        "mask": (540, 420, 360, 170, "动作掩码", "过滤满载或不可接收订单的站台"),
        "reward": (1000, 420, 360, 170, "奖励函数", "压缩完工时间、均衡负载、降低拥堵和站台占用"),
        "ppo": (80, 710, 360, 170, "PPO策略网络", "基于状态观测学习站台分配策略"),
        "train": (540, 710, 360, 170, "策略训练", "采样轨迹、计算优势、更新策略与价值网络"),
        "model": (1000, 710, 360, 170, "模型固化", "保存训练检查点和最终PPO模型"),
        "infer": (540, 990, 360, 150, "在线/离线推理", "加载模型，输入当前状态，输出站台分配动作"),
        "eval": (1000, 990, 360, 150, "策略评估", "统计完工时间、站台负载和策略对比指标"),
    }
    for key, (x, y, w, h, t, b) in boxes.items():
        fill = "#E8EEF5" if key in {"data", "clean", "state"} else "#F4F6F9"
        draw_box(draw, x, y, w, h, t, b, fill, "#9FB3C8")

    arrow(draw, (440, 215), (540, 215))
    arrow(draw, (900, 215), (1000, 215))
    arrow(draw, (1180, 300), (260, 420))
    arrow(draw, (440, 505), (540, 505))
    arrow(draw, (900, 505), (1000, 505))
    arrow(draw, (1180, 590), (260, 710))
    arrow(draw, (440, 795), (540, 795))
    arrow(draw, (900, 795), (1000, 795))
    arrow(draw, (1180, 880), (720, 990))
    arrow(draw, (900, 1065), (1000, 1065))
    arrow(draw, (720, 990), (260, 880))

    note_font = find_font(18)
    note = "说明：本图仅表达强化学习调度算法本身，不展开整体仿真过程，也不纳入近期库存约束改造逻辑。"
    draw.text((80, 1140), note, font=note_font, fill="#555555")
    img.save(DIAGRAM_PATH)


def build_doc():
    create_diagram()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = "微软雅黑"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size, color in [
        ("Heading 1", 16, "1F4D78"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 11, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("潍柴订单分拣仿真与调度算法\n过程留存记录及算法框架图")
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(16)
    srun = subtitle.add_run("适用口径：原始订单分拣仿真与调度算法，不纳入近期库存约束改造内容")
    srun.font.name = "微软雅黑"
    srun._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    srun.font.size = Pt(10.5)
    srun.font.color.rgb = RGBColor.from_string("555555")

    add_kv_table(
        doc,
        [
            ("项目名称", "潍柴订单分拣智能调度与仿真模型项目"),
            ("记录范围", "数据属性验证、模型试验、仿真逻辑、算法框架图"),
            ("数据口径", "以历史拣选数据、SKU平均分拣时间、订单清洗聚合、调度策略和仿真环境为主；不纳入库存快照及数据字典。"),
            ("形成依据", "项目目录中的 raw_data、backend、scenarios/order_picking、core_engine/rules、output、test 等文件与结果。"),
            ("示例日期", "2026-04-11，用于说明历史回放与仿真复现对比。"),
        ],
    )

    add_heading(doc, "一、数据属性验证过程留存记录", 1)
    add_body(
        doc,
        "数据属性验证用于确认历史拣选数据、SKU分拣时间标准和仿真配置参数能够支撑订单分拣调度模型运行。"
        "本记录按原始项目口径整理，不引用库存快照数据，也不使用“异常样本、边界样本”作为验证依据。"
    )
    add_matrix_table(
        doc,
        ["验证对象", "对应文件/模块", "验证重点", "形成结论"],
        [
            ["历史拣选明细", "raw_data/DMS拣选20260201-0429.XLSX", "检查订单号、开始时间、结束时间、确认状态、SKU、已拣选数量等关键字段。", "历史数据可用于订单聚合、分拣时间提取和按日期回放。"],
            ["SKU平均分拣时间", "backend/sku_avg_time.py；backend/DMS拣选20260201-0429_SKU平均分拣时间.txt", "按SKU计算加权平均单件分拣时间，过滤数量无效、耗时无效和确认状态不符合要求的记录。", "形成可被模型读取的标准处理时间，用于仿真中的作业时长计算。"],
            ["订单加载与聚合", "scenarios/order_picking/rl_environment.py", "读取历史数据，按订单聚合SKU、数量和处理时间，并按时间顺序排序。", "订单数据能够转换为模型环境可消费的订单列表。"],
            ["仿真配置参数", "scenarios/order_picking/config.py", "检查站台数量、传送带速度、站台容量、出库距离、截止时间等基础参数。", "参数完整，可支撑分拣线运行约束和时间计算。"],
            ["策略规则接口", "core_engine/rules/dispatch_rules.py", "检查AI策略、轮询、随机、最少负载、SPT、FIFO等策略接口。", "策略接口能够接收状态和合法动作约束，并返回站台分配结果。"],
        ],
        widths=[3.1, 4.2, 5.2, 3.6],
    )
    add_body(doc, "数据属性验证的主要步骤如下：")
    for item in [
        "读取历史拣选Excel，确认原始记录能够被 pandas 正常解析。",
        "将开始时间和结束时间转换为时间类型，计算每条拣选记录的实际耗时。",
        "剔除已拣选数量小于等于0、耗时小于等于0、确认状态不符合要求的记录。",
        "基于单件耗时进行IQR过滤，降低极端干扰记录对SKU标准时间的影响。",
        "按SKU聚合总耗时和总数量，得到加权平均单件分拣时间。",
        "将订单记录按订单号聚合为模型订单对象，形成包含SKU、数量、处理时间的仿真输入。",
        "检查站台、传送带、支线距离和截止时间等配置项是否可被环境正常加载。",
    ]:
        add_number(doc, item)
    add_body(
        doc,
        "验证结论：现有历史拣选数据经清洗后能够形成SKU标准分拣时间，订单明细能够被聚合为仿真输入，"
        "核心配置参数能够支持订单分拣仿真和调度策略试验。"
    )

    add_heading(doc, "二、订单顺序预处理逻辑", 1)
    add_body(
        doc,
        "订单顺序预处理位于历史订单聚合之后、强化学习环境接收订单之前，目的是在保留原始订单时序的基础上，"
        "对可处理订单、缺料订单、稀缺SKU订单和相邻冲突订单进行识别与轻量重排，使模型输入更贴近实际分拣执行约束。"
    )
    add_matrix_table(
        doc,
        ["处理步骤", "核心逻辑", "输出结果"],
        [
            ["基础排序", "按照 order_time、sequence_no 或 original_sequence 对订单排序，保留历史订单进入系统的基本时序。", "形成初始订单序列，并给订单补充 original_sequence。"],
            ["需求汇总", "遍历订单内 boxes，按SKU累计需求数量，并统计同一SKU涉及的订单数量。", "形成 SKU 日需求量和订单覆盖频次。"],
            ["库存视图构建", "若传入库存快照，则按SKU统计可用数量、可用库存单元数量和库存单元释放时间。若未传入库存，则跳过库存感知重排。", "得到 qty_by_sku、unit_count_by_sku、units_by_sku。"],
            ["短缺识别", "当订单所需SKU数量超过可用数量，或没有可用库存单元时，将订单标记为 initial_inventory_shortage。", "输出 shortage_orders；若 exclude_shortage=True，则不进入可处理订单序列。"],
            ["稀缺SKU识别", "根据单订单需求占比、当日需求占比、库存单元竞争关系和当日总需求是否超过可用量识别稀缺SKU。", "给订单写入 scarce_skus 和 scarcity_score。"],
            ["窗口重排", "在窗口范围内优先选择等待时间更小、短缺更少、库存单元更宽松且原始顺序更靠前的订单。", "减少相邻稀缺SKU冲突和库存等待。"],
            ["统计输出", "统计输入订单数、可处理订单数、短缺订单数、重排数量、稀缺SKU数量、重排前后等待估计和相邻冲突数量。", "形成 preprocess_stats，供后续试验记录和结果解释使用。"],
        ],
        widths=[3.0, 8.0, 5.0],
    )
    add_body(
        doc,
        "该逻辑的关键点是：没有库存快照时，订单保持按时间/序号排序的原始顺序；启用库存快照时，才会根据库存可用性、"
        "稀缺SKU和窗口内候选订单评分进行保守重排。重排不是重新生成订单，而是在有限窗口内调整处理顺序，降低连续处理同一稀缺SKU带来的等待和冲突。"
    )

    add_heading(doc, "三、sku_avg_time筛选逻辑及与上一版区别", 1)
    add_body(
        doc,
        "sku_avg_time用于从历史拣选明细中提取SKU标准单件分拣时间，是模型计算订单处理时长的基础。"
        "当前版筛选逻辑重点是保留整装和拆零场景下的全部正常耗时，不再将极短但真实的整装拣选耗时作为异常值剔除。"
    )
    add_matrix_table(
        doc,
        ["当前版处理步骤", "筛选/计算逻辑", "说明"],
        [
            ["读取历史数据", "读取 raw_data/DMS拣选20260201-0429.XLSX 第一张表。", "作为SKU分拣时间提取的数据源。"],
            ["时间解析", "将开始时间、结束时间转换为时间类型，计算每行耗时秒数。", "行耗时 = 结束时间 - 开始时间。"],
            ["基础有效性筛选", "保留已拣选数量 > 0、行耗时 > 0、确认代码为“确定”的记录。", "排除未实际拣选、时间无效和非确认完成记录。"],
            ["单件耗时计算", "单件耗时 = 行耗时 / 已拣选数量。", "用于衡量该SKU在单件维度上的标准作业时间。"],
            ["保留短耗时记录", "不再使用IQR剔除极短单件耗时，保留整装/拆零全部正常耗时。", "甲方确认整装拣选可能出现极短单件耗时，该类记录属于真实工艺数据。"],
            ["SKU加权平均", "按SKU分组，使用总耗时 / 总已拣选数量计算加权平均单件分拣时间。", "避免简单平均导致大批量记录权重不足。"],
            ["结果同步", "清空旧PartMaster数据后，将新的SKU标准时间写入数据库。", "供强化学习环境和历史回放读取。"],
        ],
        widths=[3.2, 7.2, 5.4],
    )
    add_matrix_table(
        doc,
        ["对比项", "上一版逻辑", "当前版逻辑", "影响"],
        [
            ["DMS范围筛选", "曾额外按服务限定符、区域、终端ID识别DMS相关记录。", "当前主要使用基础有效性条件，不再单独增加DMS范围过滤。", "减少对字段口径的重复依赖，基于当前数据源整体提取。"],
            ["极端值处理", "使用全局IQR规则过滤单件耗时极端值。", "取消IQR过滤，保留极短整装耗时和拆零耗时。", "避免把真实整装高效率记录误删，使标准时间更贴近甲方实际工艺。"],
            ["短耗时记录", "极短单件耗时可能被视为异常并剔除。", "统计并保留单件耗时 <= 0.5秒的快速整装/短耗时记录。", "降低大批量整装SKU的标准时间高估风险。"],
            ["结果落地方式", "主要生成TXT结果文件用于查看。", "直接写入PartMaster数据库，并覆盖旧标准时间。", "模型、回放和对比试验可直接读取统一标准时间。"],
            ["业务解释", "偏向数据去噪，强调过滤异常。", "偏向工艺还原，强调保留正常整装/拆零差异。", "文档中应说明短耗时不等于异常，而是工艺场景差异。"],
        ],
        widths=[2.6, 4.3, 4.3, 4.6],
    )

    add_heading(doc, "四、模型试验过程留存记录", 1)
    add_body(
        doc,
        "模型试验用于验证调度模型和规则策略在历史订单场景下能否完成工位分配、约束校验和结果输出。"
        "本项目的试验对象包括强化学习PPO策略、传统规则策略以及SimPy物理过程复现。"
    )
    add_matrix_table(
        doc,
        ["试验环节", "涉及文件/结果", "说明"],
        [
            ["模型加载", "output/models/*.zip", "加载已训练的PPO调度模型，用于逐单输出分拣站台动作。"],
            ["环境试运行", "scenarios/order_picking/rl_environment.py", "基于历史订单构建状态观测，执行动作掩码和站台状态更新。"],
            ["规则对比", "core_engine/rules/dispatch_rules.py；output/comparison_daily_results.json", "对比AI策略与Round Robin、Random等规则策略在不同启用站台数下的完工时间。"],
            ["历史回放", "backend/replay_history.py；output/history_replay_results.json", "按真实历史站台启用情况复现当日处理规模和历史完工时间。"],
            ["物理校验", "scenarios/order_picking/simpy_verify.py", "使用SimPy过程模拟订单进入工位、运输、等待和处理过程。"],
        ],
        widths=[3.2, 5.2, 7.6],
    )

    add_heading(doc, "2026-04-11复现对比记录", 2)
    add_matrix_table(
        doc,
        ["指标", "历史回放/真实口径", "仿真试验口径"],
        [
            ["日期", "2026-04-11", "2026-04-11"],
            ["有效明细行数", history_0411["valid_rows"], "策略扫描结果按聚合订单口径统计"],
            ["订单数", comparison_0411_14["orders"], comparison_0411_14["orders"]],
            ["SKU数量", history_0411["sku_count"], "使用SKU平均分拣时间参与处理时长计算"],
            ["启用站台数量", history_0411["historical_station_count"], comparison_0411_14["station_limit"]],
            ["启用站台编号", ", ".join(map(str, history_0411["historical_station_ids"])), "仿真复现按14个站台规模运行"],
            ["完工时间", seconds_to_hms(history_0411["historical_replay_makespan_seconds"]), "Round Robin: " + seconds_to_hms(comparison_0411_14["round_robin_makespan"]) + "\nAI策略: " + seconds_to_hms(comparison_0411_14["ai_makespan"])],
            ["截止时间", seconds_to_hms(history_0411["deadline_seconds"]), seconds_to_hms(history_0411["deadline_seconds"])],
        ],
        widths=[3.2, 6.3, 6.3],
    )
    add_body(
        doc,
        "从4月11日复现结果看，历史真实启用站台数量为14个，仿真试验在14个站台规模下进行复现，因此站台数量口径是一致的。"
        "完工时间不完全一致，主要原因不是站台数量差异，而是仿真模型与现场真实执行之间存在口径差别。"
    )
    add_body(
        doc,
        "本表中的订单数统一按“拣选列表/模型订单”口径统计，即一个拣选列表作为一个模型订单或调度任务；"
        "原始数据中的拣选订单字段属于上层单据口径，不作为本次订单数统计依据。"
    )
    for item in [
        "历史回放使用真实任务分布和现场执行记录，包含人工节奏、临时等待、现场组织方式等实际影响。",
        "仿真试验将订单聚合为模型可处理对象，并使用SKU平均分拣时间计算处理时长，属于标准化时间口径。",
        "AI策略和规则策略会重新给出站台分配结果，因此与历史现场分配路径不完全一致。",
        "仿真过程主要刻画传送、处理、站台容量和等待等核心逻辑，未完整还原现场所有人为干预、设备短暂停顿和作业切换损耗。",
        "因此，复现试验的意义是验证同等站台规模下模型流程可运行、约束可生效、结果可比较，而不是要求每一秒与现场记录完全一致。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "五、仿真逻辑过程留存记录", 1)
    add_body(
        doc,
        "仿真逻辑是本项目从数据输入到结果输出的完整过程。它不是额外补充的测试逻辑，而是订单分拣数字孪生模型的运行主线。"
    )
    add_body(doc, "整体仿真过程如下：")
    for item in [
        "数据加载：读取历史拣选明细和SKU平均分拣时间，形成订单处理基础数据。",
        "订单聚合：按订单号整合多个SKU明细，计算订单包含的箱/件、SKU结构和处理时间。",
        "环境初始化：设置分拣站数量、站台容量、传送带速度、出库距离、截止时间等参数。",
        "状态生成：模型环境生成当前站台负载、可用性、订单处理需求等状态观测。",
        "策略决策：AI模型或规则策略根据当前状态选择目标分拣站。",
        "合法性校验：根据站台容量和可用状态进行动作掩码，避免订单进入不可用站台。",
        "过程推进：计算订单到站运输时间、工位处理时间、等待时间和完成时间。",
        "状态更新：更新站台负载、订单完成状态和全局时间。",
        "结果输出：输出每日完工时间、站台负载、策略对比指标和可视化图表。",
    ]:
        add_number(doc, item)
    add_matrix_table(
        doc,
        ["约束类型", "项目参数/逻辑", "作用"],
        [
            ["站台数量", "Config.NUM_STATIONS = 16", "定义可参与分拣的最大站台规模。"],
            ["站台容量", "MAX_ORDERS_PER_STATION、MAX_BOXES_PER_STATION", "限制单站同时堆积订单和箱数，避免无限派发。"],
            ["传送速度", "BELT_SPEED", "用于计算订单在输送线上的移动时间。"],
            ["距离参数", "STATION_EXIT_FAR_DISTANCES、EXIT_PORT_DELTA、BRANCH_OUT_LENGTH等", "用于估计出库路径和运输时间。"],
            ["时间目标", "DEADLINE_SECONDS", "用于判断当日任务是否满足完工时限。"],
            ["策略接口", "PPO、Round Robin、Random、Least Load、SPT、FIFO", "支持AI策略和传统规则策略的对比试验。"],
        ],
        widths=[3.3, 5.1, 7.6],
    )
    add_body(
        doc,
        "仿真输出包括每日订单处理规模、启用站台数、策略完工时间、站台任务负载和整体策略对比结果。"
        "这些结果能够支撑模型效果说明、方案比选和后续算法调优。"
    )

    add_heading(doc, "六、算法框架图及说明", 1)
    doc.add_picture(str(DIAGRAM_PATH), width=Inches(6.45))
    cap = doc.add_paragraph("图1 订单分拣强化学习调度算法框架图")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("555555")

    add_body(
        doc,
        "算法框架图聚焦强化学习调度模型本身，描述从历史订单样本构建、状态空间和动作空间定义、动作掩码、奖励函数、PPO策略训练到模型固化和策略评估的过程。"
        "图中不展开整体仿真流程，仅保留强化学习算法所需的输入、学习和输出环节。"
    )
    add_matrix_table(
        doc,
        ["层级", "主要内容", "项目对应"],
        [
            ["样本输入层", "历史订单明细、SKU平均分拣时间、工位基础参数", "raw_data、backend/sku_avg_time.py、config.py"],
            ["状态空间层", "站台负载、队列占用、剩余作业量、当前订单特征", "scenarios/order_picking/rl_environment.py"],
            ["动作空间层", "在16个分拣站中选择目标站台，并通过动作掩码屏蔽不可用站台", "action_space、action_masks"],
            ["奖励函数层", "以完工时间、负载均衡、拥堵控制和站台占用成本作为优化方向", "rl_environment.py 中 step/reward 逻辑"],
            ["策略学习层", "PPO策略网络基于采样轨迹更新策略与价值网络", "train_agent_v1.py、checkpoints_v*"],
            ["模型输出层", "固化模型文件，并在试验阶段输出站台分配动作和策略评估指标", "output/models、comparison_daily_results.json"],
        ],
        widths=[3.2, 6.2, 6.4],
    )

    add_heading(doc, "七、留存结论", 1)
    for item in [
        "数据属性验证表明：历史拣选数据经过清洗、耗时计算和SKU聚合后，可以支撑订单分拣仿真输入。",
        "订单顺序预处理表明：在保留历史时序的基础上，系统可根据库存可用性和稀缺SKU冲突进行有限窗口重排，输出可追溯的预处理统计。",
        "SKU平均分拣时间提取表明：当前版保留整装和拆零正常短耗时，不再使用IQR剔除极短单件耗时，标准时间更贴近甲方实际工艺。",
        "模型试验表明：AI策略和传统规则策略均可在统一站台规模和统一约束下完成试运行，并输出可比较的完工时间结果。",
        "仿真逻辑表明：项目已形成从数据读取、订单聚合、策略决策、动作校验、时间推进到结果输出的完整仿真链路。",
        "4月11日复现说明：真实启用站台数量与仿真站台规模可以保持一致，完工时间差异主要来自调度策略、数据聚合口径、标准作业时间和现场干扰建模程度差异。",
    ]:
        add_bullet(doc, item)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("潍柴订单分拣仿真与调度算法过程留存记录")
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string("777777")

    doc.save(OUT_DOCX)


if __name__ == "__main__":
    build_doc()
    print(OUT_DOCX)
