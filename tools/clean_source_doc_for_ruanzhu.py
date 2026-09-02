from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


SRC = Path(
    r"C:\Users\14010\AppData\Roaming\Tencent\WeChatAppStore\WeChatAppStore Files"
    r"\wzhwfmwt\Files\潍柴订单分拣智能调度与数字孪生系统V1.0_源代码文档.docx"
)
OUT = Path(r"D:\weichai\weichai_model_rules_malfunction\output\docs\潍柴订单分拣智能调度与数字孪生系统V1.0_源代码文档_软著上传版.docx")


COMMENT_PREFIXES = ("//", "/*", "*/", "<!--", "-->")
COMMENT_WORDS = (
    "文件路径",
    "源代码",
    "说明",
    "模块",
    "开始",
    "结束",
    "版权",
    "作者",
    "TODO",
    "FIXME",
    "NOTE",
)
CODE_KEYWORDS = (
    "import ",
    "from ",
    "def ",
    "class ",
    "return",
    "yield",
    "if ",
    "elif ",
    "else",
    "for ",
    "while ",
    "try",
    "except",
    "finally",
    "with ",
    "async ",
    "await ",
    "raise ",
    "assert ",
    "app.",
    "db.",
    "self.",
    "@",
    "const ",
    "let ",
    "var ",
    "function ",
    "export ",
    "interface ",
    "type ",
    "public ",
    "private ",
    "protected ",
    "<",
    "</",
    "{",
    "}",
    "[",
    "]",
)
OPERATOR_HINTS = ("=", "(", ")", ",", "=>", "{", "}", "[", "]")
PATH_OR_HEADING_RE = re.compile(
    r"(^|[\s：:])([A-Za-z0-9_\-.]+/)+[A-Za-z0-9_\-.]+\.(py|js|ts|vue|json|toml|md|css|html)$"
)
FILENAME_ONLY_RE = re.compile(r"^[A-Za-z0-9_\-.]+\.(py|js|ts|vue|json|toml|md|css|html)$")


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
    runs = paragraph.runs
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def is_separator(text: str) -> bool:
    compact = text.strip()
    return len(compact) >= 3 and all(ch in "*-=#_/\\·. " for ch in compact)


def looks_like_code(text: str) -> bool:
    stripped = text.lstrip()
    if not stripped:
        return False
    if stripped.startswith(CODE_KEYWORDS):
        return True
    if re.match(r"^[A-Za-z_][A-Za-z0-9_]*\s*=", stripped):
        return True
    if re.match(r"^[A-Za-z_][A-Za-z0-9_.]*\(", stripped):
        return True
    if any(hint in stripped for hint in OPERATOR_HINTS):
        return True
    return False


def is_probable_heading_or_path(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    if any(word in stripped for word in COMMENT_WORDS):
        return True
    if any(word in stripped for word in ("入口", "文件", "路径", "模块", "说明", "源代码", "服务入口")):
        return True
    if PATH_OR_HEADING_RE.search(stripped):
        return True
    if FILENAME_ONLY_RE.match(stripped):
        return True
    if re.match(r"^[\u4e00-\u9fffA-Za-z0-9_\- ]+[：:]\s*[A-Za-z0-9_\-./\\]+\.(py|js|ts|vue|json|toml|md|css|html)$", stripped):
        return True
    return False


def is_comment_only(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return True
    if is_separator(stripped):
        return True
    if stripped.startswith(COMMENT_PREFIXES):
        return True
    if stripped in ('"""', "'''"):
        return True
    if any(word in stripped for word in COMMENT_WORDS):
        return True
    return False


def strip_inline_comment(text: str, marker: str) -> str:
    in_single = False
    in_double = False
    escape = False
    i = 0
    while i < len(text):
        ch = text[i]
        if escape:
            escape = False
            i += 1
            continue
        if ch == "\\":
            escape = True
            i += 1
            continue
        if ch == "'" and not in_double:
            in_single = not in_single
        elif ch == '"' and not in_single:
            in_double = not in_double
        if not in_single and not in_double and text.startswith(marker, i):
            if marker == "#" and i > 0 and not text[i - 1].isspace():
                i += 1
                continue
            if marker == "//" and i > 0 and text[i - 1] == ":":
                i += 2
                continue
            return text[:i].rstrip()
        i += 1
    return text.rstrip()


def clean_line(raw: str, in_triple_block: bool) -> tuple[str | None, bool, str]:
    text = raw.replace("\u00a0", " ").rstrip()
    stripped = text.strip()
    if not stripped:
        return None, in_triple_block, "blank"

    triple_starts = ('"""', "'''")
    if in_triple_block:
        if any(q in stripped for q in triple_starts):
            return None, False, "comment"
        return None, True, "comment"
    if stripped.startswith(triple_starts):
        if stripped.count('"""') == 1 or stripped.count("'''") == 1:
            return None, True, "comment"
        return None, False, "comment"

    hash_match = re.match(r"^(\s*)#\s?(.*)$", text)
    if hash_match:
        indent, after = hash_match.groups()
        after_stripped = after.strip()
        if not after_stripped:
            return None, in_triple_block, "comment"
        if after_stripped.startswith("#"):
            return None, in_triple_block, "comment"
        if is_probable_heading_or_path(after_stripped):
            return None, in_triple_block, "comment"
        if is_comment_only(after_stripped):
            return None, in_triple_block, "comment"
        if looks_like_code(after):
            cleaned = indent + after.rstrip()
            cleaned = strip_inline_comment(cleaned, "#")
            cleaned = strip_inline_comment(cleaned, "//")
            return (cleaned if cleaned.strip() else None), in_triple_block, "uncommented"
        return None, in_triple_block, "comment"

    if is_probable_heading_or_path(stripped):
        return None, in_triple_block, "comment"

    if is_comment_only(stripped):
        return None, in_triple_block, "comment"

    cleaned = strip_inline_comment(text, "#")
    cleaned = strip_inline_comment(cleaned, "//")
    if not cleaned.strip():
        return None, in_triple_block, "comment"
    return cleaned.rstrip(), in_triple_block, "kept"


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document(str(SRC))
    stats = {
        "paragraphs": len(doc.paragraphs),
        "blank_removed": 0,
        "comment_removed": 0,
        "uncommented_code_kept": 0,
        "kept": 0,
    }
    in_triple_block = False
    for paragraph in list(doc.paragraphs):
        cleaned, in_triple_block, action = clean_line(paragraph.text, in_triple_block)
        if action == "blank":
            stats["blank_removed"] += 1
        elif action == "comment":
            stats["comment_removed"] += 1
        elif action == "uncommented":
            stats["uncommented_code_kept"] += 1
            stats["kept"] += 1
        elif action == "kept":
            stats["kept"] += 1

        if cleaned is None:
            delete_paragraph(paragraph)
        else:
            replace_paragraph_text(paragraph, cleaned)

    doc.save(str(OUT))
    print(OUT)
    print(stats)


if __name__ == "__main__":
    main()
