from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output" / "docs"
OUT_PATH = OUT_DIR / "订单分拣仿真与智能调度系统用户操作手册.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
GRAY = RGBColor(90, 96, 105)
LIGHT_BLUE_GRAY = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
BORDER = "C9D3DF"


def set_run_font(run, size=None, bold=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_in):
    table.autofit = False
    table.allow_autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = sum(int(w * 1440) for w in widths_in)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total))

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths_in:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(w * 1440)))
        grid.append(col)

    for row in table.rows:
        for idx, width in enumerate(widths_in):
            if idx >= len(row.cells):
                continue
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def add_paragraph(doc, text="", style=None, bold=False, color=None, size=None, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    set_run_font(run, size={1: 16, 2: 13, 3: 12}[level], bold=True, color=BLUE if level < 3 else DARK_BLUE)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_run_font(run)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_BLUE_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], header_fill)
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_run_font(r, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run_font(r)
    set_table_geometry(table, widths)
    doc.add_paragraph()
    return table


def add_callout(doc, title, body, fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    set_table_borders(table, color="D7DEE8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, bold=True, color=NAVY)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2)
    doc.add_paragraph()


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.bold = True
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = BLUE
    styles["Heading 1"].paragraph_format.space_before = Pt(18)
    styles["Heading 1"].paragraph_format.space_after = Pt(10)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = BLUE
    styles["Heading 2"].paragraph_format.space_before = Pt(14)
    styles["Heading 2"].paragraph_format.space_after = Pt(7)
    styles["Heading 3"].font.size = Pt(12)
    styles["Heading 3"].font.color.rgb = DARK_BLUE
    styles["Heading 3"].paragraph_format.space_before = Pt(10)
    styles["Heading 3"].paragraph_format.space_after = Pt(5)

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("订单分拣仿真与智能调度系统用户操作手册")
    set_run_font(r, size=9, color=GRAY)

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("内部交付资料 | 当前版本为原型系统操作说明")
    set_run_font(r, size=9, color=GRAY)


def build_doc():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("用户操作手册")
    set_run_font(r, size=14, bold=True, color=GRAY)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run("订单分拣仿真与智能调度系统")
    set_run_font(r, size=26, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    r = subtitle.add_run("适用于当前项目原型系统的部署、启动、仿真推演与结果查看")
    set_run_font(r, size=12.5, color=GRAY)

    add_table(
        doc,
        ["项目", "说明"],
        [
            ["文档性质", "用户操作手册（初版）"],
            ["系统定位", "订单分拣仿真与智能调度原型系统"],
            ["适用对象", "甲方业务人员、现场演示人员、实施与技术支持人员"],
            ["覆盖范围", "环境准备、服务启动、数据准备、仿真运行、结果查看、常见问题"],
            ["当前限制", "系统尚未形成安装包或商业软件完整产品化交付形态"],
        ],
        [1.6, 4.9],
    )

    add_callout(
        doc,
        "交付口径建议",
        "本文档按“可操作的软件原型系统”编写。当前项目已经具备前端界面、后端接口、仿真引擎、调度模型和结果展示能力，但仍建议在正式沟通中称为“原型系统”或“仿真系统”，避免承诺尚未实现的产品化能力。",
    )

    add_heading(doc, "1. 系统概述", 1)
    add_paragraph(
        doc,
        "订单分拣仿真与智能调度系统用于对订单分拣任务进行数据导入、库存预处理、智能调度策略计算、离散事件仿真和可视化展示。当前项目包含 FastAPI 后端、Vue 3 前端大屏、强化学习调度模型、SimPy 仿真逻辑、库存与订单预处理脚本，以及仿真结果输出文件。",
    )
    add_table(
        doc,
        ["模块", "当前项目中的作用"],
        [
            ["前端大屏", "提供批次输入、库存快照选择、仿真启动、进度查看、3D 数字孪生动画和算法对比展示。"],
            ["后端服务", "提供订单上传、仿真启动、任务状态查询、playbook 获取、库存快照查询等 API。"],
            ["仿真引擎", "基于站台、传送距离、处理时间、订单约束等因素执行离散事件仿真。"],
            ["调度模型", "使用强化学习模型或规则策略生成订单到站台的分配方案。"],
            ["数据处理", "支持历史订单、库存快照、SKU 平均拣选时间等数据准备与预处理。"],
            ["结果输出", "输出调度结果、仿真记录、对比汇总、3D 动画脚本等文件。"],
        ],
        [1.45, 5.05],
    )

    add_heading(doc, "2. 使用前准备", 1)
    add_heading(doc, "2.1 运行环境", 2)
    add_table(
        doc,
        ["类别", "要求"],
        [
            ["操作系统", "Windows 环境，当前项目路径示例为 D:\\weichai\\weichai_model_rules_malfunction。"],
            ["Python", "用于后端服务、仿真引擎、数据处理和模型加载。项目内已有 requirements.txt。"],
            ["Node.js", "用于启动 Vue 3 + Vite 前端。前端目录包含 package.json。"],
            ["数据库", "当前 README 中描述为 MySQL，默认库名 weichai_aps。"],
            ["浏览器", "建议使用 Chrome、Edge 等现代浏览器访问前端页面。"],
        ],
        [1.3, 5.2],
    )
    add_heading(doc, "2.2 项目目录说明", 2)
    add_table(
        doc,
        ["目录或文件", "说明"],
        [
            ["backend/", "FastAPI 服务、数据库模型、订单导入、SKU 平均耗时计算等。"],
            ["weichai-aps-frontend/", "Vue 3 前端大屏与 3D 展示界面。"],
            ["scenarios/order_picking/", "订单分拣仿真场景、强化学习环境、训练和对比脚本。"],
            ["core_engine/", "资源模型、规则模型、传送与站台处理逻辑。"],
            ["data/", "轻量输入数据、库存快照、历史订单 JSON。"],
            ["raw_data/", "原始 Excel 数据。"],
            ["output/", "模型、仿真结果、对比结果、playbook 与动画脚本。"],
        ],
        [1.8, 4.7],
    )

    add_heading(doc, "3. 系统启动", 1)
    add_heading(doc, "3.1 启动后端服务", 2)
    add_paragraph(doc, "在 PowerShell 中进入项目根目录并激活 Python 环境，然后启动后端服务。")
    add_numbered(
        doc,
        [
            "进入项目目录：cd D:\\weichai\\weichai_model_rules_malfunction",
            "激活环境：.\\venv\\Scripts\\activate",
            "进入后端目录：cd backend",
            "启动服务：python server.py",
        ],
    )
    add_callout(
        doc,
        "后端端口",
        "前端接口默认通过 /api/v1 访问，并由 Vite 代理转发到后端。项目代码中后端服务常用端口为 8088。若启动失败，应优先检查端口占用、数据库连接和 Python 依赖。",
        fill="F7FAFF",
    )

    add_heading(doc, "3.2 启动前端页面", 2)
    add_numbered(
        doc,
        [
            "新开一个 PowerShell 窗口，进入前端目录：cd D:\\weichai\\weichai_model_rules_malfunction\\weichai-aps-frontend",
            "首次运行或依赖变化后执行：npm install",
            "启动前端：npm run dev",
            "在浏览器访问 Vite 提示的地址，通常为 http://127.0.0.1:5173。",
        ],
    )

    add_heading(doc, "4. 数据准备", 1)
    add_heading(doc, "4.1 订单批次", 2)
    add_paragraph(
        doc,
        "系统以订单批次号作为一次仿真推演的主要入口。前端默认示例批次为 ORDER_WAVE_2026-04-11，实际使用时应根据已导入数据库的订单批次填写。",
    )
    add_table(
        doc,
        ["字段", "说明", "示例"],
        [
            ["batch_no", "订单批次号，用于关联同一批订单。", "ORDER_WAVE_2026-04-11"],
            ["order_id", "订单号或拣选单号。", "PICK001"],
            ["part_type", "SKU 编码，建议保留完整编码及后缀。", "610800050009:D00"],
            ["quantity", "该订单所需 SKU 数量。", "10"],
        ],
        [1.2, 3.5, 1.8],
    )
    add_heading(doc, "4.2 库存快照", 2)
    add_paragraph(
        doc,
        "当前前端支持选择日初库存快照和日末校验快照。系统默认快照示例为 2025-07-01-morning 和 2025-07-01-evening。若快照列表为空，应检查 data/inventory/ 下是否存在对应 JSON 文件，或运行库存预处理脚本生成快照。",
    )
    add_heading(doc, "4.3 SKU 平均拣选时间", 2)
    add_paragraph(
        doc,
        "SKU 平均拣选时间用于估算订单处理耗时。历史拣选数据更新后，可通过 backend/sku_avg_time.py 重新计算并同步到基础数据表。",
    )

    add_heading(doc, "5. 仿真操作流程", 1)
    add_heading(doc, "5.1 启动一次仿真", 2)
    add_numbered(
        doc,
        [
            "确认后端服务和前端页面均已启动。",
            "在前端页面进入“3D 数字孪生”视图。",
            "在订单批次输入框中填写需要推演的 batch_no。",
            "选择日初库存快照和日末校验快照。",
            "点击“启动”按钮。",
            "等待进度条从后端返回实时状态，仿真完成后系统自动加载 3D 播放脚本。",
        ],
    )
    add_heading(doc, "5.2 查看运行状态", 2)
    add_bullets(
        doc,
        [
            "进度条显示后端仿真任务当前完成比例。",
            "状态文字显示数据加载、模型计算、仿真推演、结果生成等阶段信息。",
            "若状态显示失败，应记录订单批次、库存快照、后台报错信息和操作时间，交由技术人员排查。",
        ],
    )
    add_heading(doc, "5.3 查看 3D 数字孪生", 2)
    add_paragraph(
        doc,
        "仿真完成后，前端会请求 /simulation/playbook/{task_id} 获取 3D 播放数据。3D 视图用于展示订单实体在站台、传送与处理过程中的状态变化，并联动更新 KPI 数据与站台状态矩阵。",
    )
    add_table(
        doc,
        ["显示区域", "含义"],
        [
            ["订单进度", "显示主订单完成数量、总订单数量和整体完成比例。"],
            ["箱体统计", "显示总投放实体箱、在途/排队箱体、已完成箱体等。"],
            ["站台矩阵", "显示 16 个站台的忙闲状态、订单数量和节能休眠状态。"],
            ["耗时指标", "显示包含回库等环节在内的总耗时估计。"],
        ],
        [1.5, 5.0],
    )
    add_heading(doc, "5.4 查看算法对比", 2)
    add_paragraph(
        doc,
        "前端包含“智能算法对比”页签，可用于查看 AI 策略、轮询策略、随机策略等方案在完成时间、站台利用和订单处理方面的差异。具体可见内容取决于后端本次仿真返回的数据字段。",
    )

    add_heading(doc, "6. 结果文件与输出说明", 1)
    add_table(
        doc,
        ["输出位置", "说明"],
        [
            ["output/schedule_results/", "调度结果 JSON，包含任务编号、策略、站台分配和统计信息。"],
            ["output/history_replay_*.json", "历史订单回放结果。"],
            ["output/comparison_daily_results.json", "不同策略或不同日期的对比明细。"],
            ["output/comparison_daily_summary.json", "对比结果汇总。"],
            ["output/playbooks/", "前端 3D 动画或播放脚本相关数据。"],
            ["output/models/", "强化学习模型文件，例如 ppo_masking_model_v6.zip。"],
        ],
        [2.2, 4.3],
    )
    add_callout(
        doc,
        "结果解释原则",
        "仿真输出用于评估不同调度策略在给定数据、库存快照和模型参数下的表现。结果不应直接等同于真实现场最终产能，正式使用前应结合现场设备参数、人员配置、异常停机、数据完整性进行校准。",
    )

    add_heading(doc, "7. 常见问题处理", 1)
    add_table(
        doc,
        ["问题", "可能原因", "处理建议"],
        [
            ["前端无法启动", "Node 依赖未安装或端口被占用。", "执行 npm install；检查 Vite 输出端口；必要时更换端口。"],
            ["点击启动后无响应", "后端未启动或代理未转发成功。", "确认 backend/server.py 正在运行；检查 8088 端口。"],
            ["提示批次无订单", "数据库中不存在该 batch_no。", "确认订单已导入；检查批次号拼写。"],
            ["库存快照为空", "data/inventory 下缺少快照文件。", "运行库存预处理脚本或放入对应 JSON 快照。"],
            ["模型加载失败", "output/models 下缺少 .zip 模型文件。", "确认模型文件存在，或重新训练/复制模型。"],
            ["页面文字乱码", "源文件或浏览器编码/字体设置异常。", "统一使用 UTF-8 保存前端和说明文档；检查中文字体支持。"],
            ["数据库连接失败", "MySQL 未启动、库未创建或账号密码不匹配。", "检查 backend/database.py 中连接配置与本机数据库状态。"],
        ],
        [1.55, 2.25, 2.7],
    )

    add_heading(doc, "8. 当前版本限制", 1)
    add_bullets(
        doc,
        [
            "当前系统更适合作为仿真与调度算法原型系统使用，尚未形成独立安装包。",
            "用户权限、审计日志、参数配置界面、异常恢复流程等产品化能力仍需进一步完善。",
            "部分数据准备仍依赖脚本和指定目录，非技术用户直接操作门槛较高。",
            "仿真结果依赖当前数据质量、模型版本、站台参数、库存快照和算法配置。",
            "正式验收时建议明确“原型系统功能范围”和“不包含项”，避免对软件成熟度产生误解。",
        ],
    )

    add_heading(doc, "9. 建议验收检查清单", 1)
    add_table(
        doc,
        ["检查项", "通过标准"],
        [
            ["后端服务", "能够正常启动，API 无启动报错。"],
            ["前端页面", "能够访问首页，页签切换正常。"],
            ["订单批次", "输入有效 batch_no 后可启动仿真。"],
            ["库存快照", "可选择日初和日末快照。"],
            ["仿真进度", "任务启动后可轮询状态并显示进度。"],
            ["仿真结果", "任务完成后可获取结果和 playbook。"],
            ["3D 展示", "仿真完成后可播放或展示订单运行过程。"],
            ["输出文件", "output 目录下生成对应结果文件。"],
        ],
        [2.0, 4.5],
    )

    add_heading(doc, "10. 附录：主要接口", 1)
    add_table(
        doc,
        ["接口", "用途"],
        [
            ["POST /api/v1/orders/upload", "上传订单批次数据。"],
            ["POST /api/v1/simulation/start", "启动仿真推演任务。"],
            ["GET /api/v1/simulation/status/{task_id}", "查询仿真任务状态。"],
            ["GET /api/v1/simulation/playbook/{task_id}", "获取 3D 展示播放数据。"],
            ["POST /api/v1/schedule/dispatch", "启动调度分配任务。"],
            ["GET /api/v1/schedule/result/{task_id}", "获取调度结果。"],
            ["GET /api/v1/inventory/snapshots", "获取库存快照列表。"],
            ["GET /api/v1/inventory/summary/{snapshot_id}", "获取库存快照摘要。"],
            ["GET /api/v1/model/training_metrics", "获取模型训练指标。"],
        ],
        [2.65, 3.85],
    )

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_heading(doc, "版本说明", 1)
    add_paragraph(
        doc,
        "本手册为基于当前项目代码和目录结构整理的初版操作说明。后续若系统完成安装包、权限体系、正式部署脚本、参数配置页面或验收版本冻结，应同步更新本文档中的环境准备、启动方式、操作截图和故障处理章节。",
    )

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_doc())
