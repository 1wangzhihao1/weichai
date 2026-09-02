from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output" / "docs"
ASSET_DIR = OUT_DIR / "manual_assets"
OUT_PATH = OUT_DIR / "潍柴APS智能排产系统软件说明书_详细版.docx"

ASSETS = {
    "frontend_home": ROOT / "work" / "ppt_edit_notes" / "tmp" / "frontend_edge.png",
    "frontend_3d": ROOT / "work" / "ppt_edit_notes" / "tmp" / "assets" / "frontend_3d_view.png",
    "frontend_compare": ROOT / "work" / "ppt_edit_notes" / "tmp" / "assets" / "frontend_compare_view.png",
    "frontend_log": ROOT / "work" / "ppt_edit_notes" / "tmp" / "assets" / "frontend_log_view.png",
    "api_docs": ASSET_DIR / "api_docs.png",
    "performance": ROOT / "output" / "performance_comparison.png",
}

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
GRAY = RGBColor(95, 103, 112)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F5F7FA"
BORDER = "C9D3DF"


def set_run_font(run, size=None, bold=None, color=None, font="Calibri"):
    run.font.name = font
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=130, bottom=90, end=130):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    table.autofit = False
    table.allow_autofit = False
    total = sum(int(w * 1440) for w in widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total))
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")
    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                cell = row.cells[idx]
                cell.width = Inches(width)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = tc_pr.first_child_found_in("w:tcW")
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:type"), "dxa")
                tc_w.set(qn("w:w"), str(int(width * 1440)))
                set_cell_margins(cell)


def paragraph(doc, text="", style=None, size=None, bold=False, color=None, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading("", level=level)
    run = p.add_run(text)
    sizes = {1: 16, 2: 13, 3: 12}
    colors = {1: BLUE, 2: BLUE, 3: DARK_BLUE}
    set_run_font(run, size=sizes[level], bold=True, color=colors[level])
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_run_font(run)
    return p


def number(doc, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    set_run_font(run)
    return p


def code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    set_table_borders(table, color="D8DEE8")
    cell = table.cell(0, 0)
    shade_cell(cell, "F3F5F8")
    cell.text = ""
    for idx, line in enumerate(lines):
        p = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        set_run_font(r, size=9.5, font="Consolas", color=NAVY)
    doc.add_paragraph()


def table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(t, widths)
    set_table_borders(t)
    for idx, header in enumerate(headers):
        cell = t.rows[0].cells[idx]
        shade_cell(cell, LIGHT_BLUE)
        cell.text = ""
        r = cell.paragraphs[0].add_run(header)
        set_run_font(r, bold=True, color=NAVY)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    for row in rows:
        cells = t.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = ""
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run_font(r)
    set_table_geometry(t, widths)
    doc.add_paragraph()
    return t


def callout(doc, title, body):
    t = doc.add_table(rows=1, cols=1)
    set_table_geometry(t, [6.5])
    set_table_borders(t, color="D8DEE8")
    cell = t.cell(0, 0)
    shade_cell(cell, LIGHT_GRAY)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, bold=True, color=NAVY)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2)
    doc.add_paragraph()


def image(doc, key, caption):
    path = ASSETS[key]
    if not path.exists():
        callout(doc, "图片占位", f"未找到图片文件：{path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.25))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_run_font(r, size=9.5, color=GRAY)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, side, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("潍柴 APS 智能排产数字孪生系统软件说明书")
    set_run_font(r, size=9, color=GRAY)

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("内部交付资料 | 初版说明书")
    set_run_font(r, size=9, color=GRAY)


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure(doc)

    paragraph(doc, "软件说明书", size=14, bold=True, color=GRAY)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run("潍柴 APS 智能排产数字孪生系统")
    set_run_font(r, size=26, bold=True, color=NAVY)
    paragraph(
        doc,
        "基于当前 README 与项目代码整理，覆盖环境准备、数据导入、后端接口、前端操作、仿真运行、结果查看与常见问题。",
        size=12.5,
        color=GRAY,
    )
    table(
        doc,
        ["项目", "说明"],
        [
            ["文档版本", "初版，用于效果预览和后续修改"],
            ["适用对象", "演示人员、实施人员、技术支持人员和甲方业务试用人员"],
            ["系统定位", "订单分拣智能调度与数字孪生仿真原型系统"],
            ["主要技术", "FastAPI、MySQL、Vue 3、Vite、Three.js、SimPy、强化学习调度模型"],
            ["访问入口", "后端接口：http://127.0.0.1:8088/docs；前端页面：http://127.0.0.1:5173"],
        ],
        [1.55, 4.95],
    )
    image(doc, "frontend_home", "图 1 前端 3D 数字孪生主界面与仿真控制区")

    heading(doc, "1. 系统概述", 1)
    paragraph(
        doc,
        "本系统用于对订单分拣任务进行数据准备、库存预处理、智能调度、离散事件仿真和可视化展示。系统先从订单与库存数据中筛选可执行订单，再使用 AI 或规则策略完成订单到站台的分配，最后通过前端 3D 场景展示站台状态、订单进度和执行结果。",
    )
    table(
        doc,
        ["模块", "作用"],
        [
            ["后端服务", "提供订单上传、批次查询、库存快照查询、SKU 工时重建、仿真启动、任务状态查询和 playbook 获取接口。"],
            ["前端大屏", "提供批次选择、库存快照选择、仿真启动、进度展示、16 站台状态矩阵、3D 数字孪生和算法对比页面。"],
            ["数据处理", "从拣选 Excel 重建 SKU 平均处理时间，从库存 Excel 生成库存快照，并在仿真前识别缺料订单。"],
            ["调度与仿真", "使用强化学习模型或规则策略进行站台分配，再由 SimPy 执行离散事件仿真。"],
            ["结果输出", "输出调度结果、仿真报告、3D 播放脚本、模型文件和策略对比数据。"],
        ],
        [1.45, 5.05],
    )

    heading(doc, "2. 目录与配置", 1)
    heading(doc, "2.1 项目目录", 2)
    table(
        doc,
        ["目录或文件", "说明"],
        [
            ["backend/", "FastAPI 服务、数据库初始化、订单导入、SKU 工时计算和接口实现。"],
            ["weichai-aps-frontend/", "Vue 3 + Vite 前端项目，包含 3D 数字孪生组件与算法对比组件。"],
            ["scenarios/order_picking/", "订单分拣场景、强化学习环境、训练脚本、库存预处理和模型筛选脚本。"],
            ["core_engine/", "仿真引擎、资源模型和调度规则。"],
            ["raw_data/", "历史拣选、每日拣选、库存等原始 Excel 数据。"],
            ["data/inventory/", "预处理后的库存快照 JSON。"],
            ["output/", "模型、playbook、调度结果、仿真报告和对比图表。"],
        ],
        [1.85, 4.65],
    )
    heading(doc, "2.2 关键配置", 2)
    paragraph(doc, "主要配置文件为 config/app_config.toml。切换日期、批次、库存快照或模型时，优先修改该文件。")
    code_block(
        doc,
        [
            '[datasets]',
            'active_date = "2025-07-01"',
            "",
            '[simulation]',
            'default_batch_no = "ORDER_WAVE_2025-07-01"',
            'default_initial_snapshot_id = "2025-07-01-morning"',
            'default_evening_snapshot_id = "2025-07-01-evening"',
            'default_strategy = "ai"',
            "",
            '[model]',
            'active_model = "ppo_masking_model_v6.zip"',
            'model_dir = "output/models"',
        ],
    )
    table(
        doc,
        ["配置段", "配置项", "含义", "什么时候修改"],
        [
            ["datasets", "active_date", "当前业务日期。程序会优先到 raw_data/daily/<日期>/ 下查找当天拣选和库存文件。", "新增一天数据或切换演示日期时修改。"],
            ["simulation", "default_batch_no", "前端默认选中的订单波次号，通常与导入订单时使用的 batch_no 保持一致。", "更换订单波次或重新导入订单后修改。"],
            ["simulation", "default_initial_snapshot_id", "默认日初库存快照 ID，一般对应 data/inventory/<日期>-morning.json。", "更换日初库存快照时修改。"],
            ["simulation", "default_evening_snapshot_id", "默认日末校验库存快照 ID，一般对应 data/inventory/<日期>-evening.json。", "更换日末校验库存快照时修改。"],
            ["simulation", "default_strategy", "默认调度策略。ai 表示使用强化学习/智能调度策略。", "需要切换 AI、规则或其他策略时修改。"],
            ["model", "active_model", "当前后端加载的模型文件名。系统会在 model_dir 目录中查找该文件。", "更换正式模型或模型版本时修改。"],
            ["model", "model_dir", "模型文件所在目录，当前为 output/models。", "模型目录调整时修改，一般保持默认。"],
        ],
        [0.95, 2.2, 2.25, 1.1],
    )
    callout(
        doc,
        "配置文件的作用",
        "这些配置只负责告诉系统默认读哪一天、哪个批次、哪个库存快照和哪个模型；它们不会自动创建数据库、导入订单或生成库存快照。新增数据后，仍需要按后续章节执行对应的数据处理或导入操作。",
    )

    heading(doc, "3. 环境准备", 1)
    heading(doc, "3.1 Python 环境", 2)
    paragraph(doc, "首次运行后端、仿真和数据处理脚本前，需要创建并激活 Python 虚拟环境。")
    code_block(
        doc,
        [
            r"cd D:\weichai\weichai_model_rules_malfunction",
            "python -m venv venv",
            r".\venv\Scripts\activate",
            "pip install -r requirements.txt",
        ],
    )
    heading(doc, "3.2 前端依赖", 2)
    paragraph(doc, "首次运行前端或 package.json 依赖发生变化后，需要在前端目录安装 Node 依赖。")
    code_block(
        doc,
        [
            r"cd D:\weichai\weichai_model_rules_malfunction\weichai-aps-frontend",
            "npm install",
        ],
    )
    heading(doc, "3.3 数据库初始化", 2)
    paragraph(doc, "项目 README 中默认使用 MySQL 数据库 weichai_aps。先创建数据库，再运行初始化脚本创建基础表。")
    paragraph(
        doc,
        "CREATE DATABASE 命令需要在 MySQL 客户端中运行，不是在 Python 终端中运行。常见运行位置包括 MySQL Command Line Client、MySQL Workbench 的 Query 标签页、Navicat/DBeaver 的 SQL 查询窗口，或 PowerShell 中登录 mysql 客户端后执行。",
    )
    code_block(
        doc,
        [
            "# 方式一：PowerShell 进入 MySQL 客户端",
            "mysql -u root -p",
            "",
            "# 登录成功后，在 mysql> 提示符后执行",
            "CREATE DATABASE weichai_aps DEFAULT CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci;",
            "",
            "# 然后回到项目终端，初始化表结构",
            r"cd D:\weichai\weichai_model_rules_malfunction\backend",
            r"..\venv\Scripts\activate",
            "python database.py",
        ],
    )
    table(
        doc,
        ["主要表", "用途"],
        [
            ["t_part_master", "SKU 标准处理时间。"],
            ["t_station_master", "站台基础信息。"],
            ["t_order_pool", "订单主表。"],
            ["t_order_bom", "订单 SKU 明细。"],
            ["t_simulation_task", "仿真任务记录。"],
            ["t_dispatch_result", "调度结果记录。"],
        ],
        [2.0, 4.5],
    )

    heading(doc, "4. 数据准备", 1)
    heading(doc, "4.1 重建 SKU 平均处理时间", 2)
    paragraph(doc, "SKU 平均处理时间由 backend/sku_avg_time.py 计算，结果写入 t_part_master。后端启动仿真时也可根据配置自动刷新。")
    code_block(
        doc,
        [
            r"cd D:\weichai\weichai_model_rules_malfunction\backend",
            r"..\venv\Scripts\activate",
            "python sku_avg_time.py",
        ],
    )
    callout(
        doc,
        "操作提示",
        "如果脚本执行失败，优先查看终端最后几行报错，并核对虚拟环境、Excel 文件路径、数据库连接和后端服务状态。",
    )
    heading(doc, "4.2 生成库存快照", 2)
    paragraph(doc, "库存 Excel 会先预处理为轻量 JSON，前端下拉框读取的也是这些快照。")
    callout(
        doc,
        "是否每次都要重新运行",
        "新增某一天的库存 Excel 后，需要至少运行一次 inventory_preprocess.py，把 Excel 转成 data/inventory/ 下的 JSON 快照。如果不运行，也没有手动放入对应 JSON，前端库存下拉框通常就找不到这一天的日初/日末库存快照。已经生成过且 Excel 未变化时，不需要每次启动前端都重新运行。",
    )
    code_block(
        doc,
        [
            r"cd D:\weichai\weichai_model_rules_malfunction",
            r".\venv\Scripts\activate",
            r"python scenarios\order_picking\inventory_preprocess.py",
        ],
    )
    paragraph(doc, "生成结果示例：data/inventory/2025-07-01-morning.json、data/inventory/2025-07-01-evening.json。")
    heading(doc, "4.3 导入订单", 2)
    paragraph(doc, "正式订单通过 POST /api/v1/orders/upload 导入。项目内置脚本可将 Excel 转为接口 JSON 后上传。")
    paragraph(
        doc,
        "导入前需要确认后端已启动、MySQL 数据库已创建并初始化。POST 导入的本质是把订单批次 JSON 发送给后端接口，由后端写入 t_order_pool 和 t_order_bom 等订单表。",
    )
    heading(doc, "4.3.1 通过接口文档页面导入", 3)
    number(doc, "启动后端：python server.py。")
    number(doc, "浏览器打开 http://127.0.0.1:8088/docs。")
    number(doc, "展开 POST /api/v1/orders/upload。")
    number(doc, "点击 Try it out，在 Request body 中粘贴订单 JSON。")
    number(doc, "点击 Execute，返回成功后即可在前端订单波次下拉框中看到对应 batch_no。")
    heading(doc, "4.3.2 通过 PowerShell 调用接口导入", 3)
    paragraph(doc, "如果甲方系统已经给出订单 JSON 文件，可以用 PowerShell 直接调用后端接口导入。下面按实际操作顺序说明。")
    callout(
        doc,
        "操作前确认",
        "先确认 MySQL 已启动、数据库表已经初始化、后端 server.py 正在运行，并且浏览器可以打开 http://127.0.0.1:8088/docs。否则 PowerShell 调接口会连接失败。",
    )
    heading(doc, "第一步：准备订单 JSON 文件", 3)
    bullet(doc, "新建一个文件夹，例如 D:\\weichai\\customer_orders。")
    bullet(doc, "在该文件夹中新建 orders_2025_07_01.json。")
    bullet(doc, "用 VS Code、记事本或其他文本编辑器打开该文件。")
    bullet(doc, "把订单 JSON 内容粘贴进去，并保存为 UTF-8 编码。")
    paragraph(doc, "订单 JSON 的结构如下。batch_no 是本次订单波次号，前端启动仿真时需要选择同一个波次。orders 数组中的每一行表示某个订单需要的一个 SKU 及数量。")
    code_block(
        doc,
        [
            "{",
            '  "batch_no": "ORDER_WAVE_2025-07-01",',
            '  "orders": [',
            '    {"order_id": "PICK001", "part_type": "610800050009:D00", "quantity": 10},',
            '    {"order_id": "PICK001", "part_type": "610800050010:D00", "quantity": 2},',
            '    {"order_id": "PICK002", "part_type": "610800050011:D00", "quantity": 1}',
            "  ]",
            "}",
        ],
    )
    table(
        doc,
        ["字段", "怎么填"],
        [
            ["batch_no", "订单波次号。建议按日期命名，例如 ORDER_WAVE_2025-07-01。前端后续也要选这个波次。"],
            ["orders", "订单明细数组。一个订单包含多个 SKU 时，可以出现多行相同 order_id。"],
            ["order_id", "订单号或拣选单号，例如 PICK001。"],
            ["part_type", "SKU 编码，要与库存和 SKU 工时数据中的编码保持一致。"],
            ["quantity", "该订单需要的该 SKU 数量，只填数字，不加单位。"],
        ],
        [1.35, 5.15],
    )
    heading(doc, "第二步：打开 PowerShell 并读取 JSON 文件", 3)
    paragraph(doc, "打开 PowerShell 后，先执行下面命令，把 JSON 文件内容读取到变量 $json 中。路径要改成自己实际保存的 JSON 文件路径。")
    code_block(
        doc,
        [
            "$json = Get-Content \"D:\\weichai\\customer_orders\\orders_2025_07_01.json\" -Raw -Encoding UTF8",
        ],
    )
    paragraph(doc, "如果想确认文件是否读到了，可以继续执行下面命令，PowerShell 会打印 JSON 文本。")
    code_block(
        doc,
        [
            "$json",
        ],
    )
    heading(doc, "第三步：调用后端上传接口", 3)
    paragraph(doc, "确认 $json 有内容后，再执行上传命令。下面是多行写法，行尾的反引号 ` 表示命令下一行继续，复制时不要漏掉。")
    code_block(
        doc,
        [
            "Invoke-RestMethod `",
            "  -Uri http://127.0.0.1:8088/api/v1/orders/upload `",
            "  -Method Post `",
            "  -ContentType \"application/json; charset=utf-8\" `",
            "  -Body $json",
        ],
    )
    paragraph(doc, "如果担心多行命令的反引号复制出错，也可以使用下面的一行写法，效果完全相同。")
    code_block(
        doc,
        [
            "Invoke-RestMethod -Uri http://127.0.0.1:8088/api/v1/orders/upload -Method Post -ContentType \"application/json; charset=utf-8\" -Body $json",
        ],
    )
    heading(doc, "第四步：判断是否导入成功", 3)
    paragraph(doc, "命令执行成功时，PowerShell 会返回后端响应结果，通常会包含 code、message、data 或导入数量等信息。只要没有红色报错，并且返回成功状态，就说明接口调用成功。")
    code_block(
        doc,
        [
            "# 查询后端已导入的订单波次",
            "Invoke-RestMethod -Uri http://127.0.0.1:8088/api/v1/orders/batches -Method Get",
        ],
    )
    paragraph(doc, "如果返回结果里能看到 ORDER_WAVE_2025-07-01，说明该批次已进入系统。随后刷新前端页面，在订单波次下拉框里选择该 batch_no 即可启动仿真。")
    heading(doc, "第五步：常见错误处理", 3)
    table(
        doc,
        ["现象", "原因", "处理方式"],
        [
            ["无法连接 127.0.0.1:8088", "后端没有启动，或端口不是 8088。", "先在 backend 目录运行 python server.py，并确认 /docs 能打开。"],
            ["Get-Content 找不到路径", "JSON 文件路径写错，或文件不存在。", "检查文件是否真的在 D:\\weichai\\customer_orders\\ 下，路径包含空格时要加英文双引号。"],
            ["JSON 解析失败", "JSON 格式不合法，例如少逗号、少引号、末尾多逗号。", "用 VS Code 检查 JSON，确保键名和字符串都用英文双引号。"],
            ["接口返回批次为空或订单数为 0", "orders 数组为空，或字段名不是 order_id、part_type、quantity。", "按示例字段名重新整理 JSON。"],
            ["前端下拉框看不到波次", "前端未刷新，或导入的 batch_no 与配置/查询结果不一致。", "刷新前端页面，并用 /api/v1/orders/batches 查询实际导入的 batch_no。"],
        ],
        [1.55, 2.2, 2.75],
    )
    heading(doc, "4.3.3 通过项目脚本模拟接口导入", 3)
    paragraph(
        doc,
        "import_july1_picking.py 用于读取每日拣选 Excel，将 Excel 中的订单行转换成接口需要的 JSON，并调用 /api/v1/orders/upload 上传。使用 --mode api 时，它不是绕过接口直接写前端，而是模拟甲方系统通过 HTTP 接口把订单导入后端，最终仍由后端写入数据库。使用 --mode db 时则是调试用的直连数据库写入方式，会绕过接口。",
    )
    code_block(
        doc,
        [
            r"cd D:\weichai\weichai_model_rules_malfunction\backend",
            r"..\venv\Scripts\activate",
            r"python import_july1_picking.py --mode api --api-url http://127.0.0.1:8088/api/v1",
        ],
    )
    table(
        doc,
        ["参数", "说明"],
        [
            ["--mode api", "按正式接口路径导入：Excel -> JSON -> POST /orders/upload -> 后端写数据库。推荐用于演示接口导入流程。"],
            ["--mode db", "脚本直接连接数据库写入订单表。适合本地调试，不适合展示“甲方通过接口导入”。"],
            ["--api-url", "后端 API 基础地址。默认本机为 http://127.0.0.1:8088/api/v1。"],
            ["--date", "指定导入哪一天的 raw_data/daily/<date>/picking 数据。未指定时使用配置中的默认日期。"],
            ["--batch-no", "手动指定订单波次号。不指定时通常生成 ORDER_WAVE_<date>。"],
            ["--dry-run", "只解析并打印统计，不真正写数据库或调用接口。用于导入前检查。"],
        ],
        [1.45, 5.05],
    )
    table(
        doc,
        ["字段", "说明", "示例"],
        [
            ["batch_no", "订单波次号，前端按该值启动仿真。", "ORDER_WAVE_2025-07-01"],
            ["order_id", "订单号或拣选单号。", "PICK001"],
            ["part_type", "完整 SKU 编码。", "610800050009:D00"],
            ["quantity", "该订单所需 SKU 数量。", "10"],
        ],
        [1.2, 3.5, 1.8],
    )

    heading(doc, "5. 启动系统", 1)
    heading(doc, "5.1 启动后端", 2)
    paragraph(doc, "后端启动后默认监听 http://127.0.0.1:8088，接口文档访问 /docs。")
    code_block(
        doc,
        [
            r"cd D:\weichai\weichai_model_rules_malfunction\backend",
            r"..\venv\Scripts\activate",
            "python server.py",
        ],
    )
    image(doc, "api_docs", "图 3 后端 FastAPI 接口文档页面")
    heading(doc, "5.2 启动前端", 2)
    paragraph(doc, "前端默认由 Vite 启动，通常访问地址为 http://127.0.0.1:5173。")
    code_block(
        doc,
        [
            r"cd D:\weichai\weichai_model_rules_malfunction\weichai-aps-frontend",
            "npm run dev",
        ],
    )
    image(doc, "frontend_home", "图 4 前端启动后的主操作界面")

    heading(doc, "6. 前端仿真操作", 1)
    heading(doc, "6.1 选择批次和库存快照", 2)
    number(doc, "确认后端 8088 和前端 5173 均已启动。")
    number(doc, "进入 3D 数字孪生页签。")
    number(doc, "在订单波次下拉框中选择或确认 batch_no。")
    number(doc, "选择日初库存快照和日末校验快照。")
    number(doc, "确认站台状态矩阵处于待命状态。")
    image(doc, "frontend_home", "图 5 订单批次、库存快照和站台状态矩阵")
    heading(doc, "6.2 启动仿真", 2)
    number(doc, "点击“启动”按钮。")
    number(doc, "前端调用 /api/v1/simulation/start。")
    number(doc, "后端返回 task_id 后，前端每秒轮询 /api/v1/simulation/status/{task_id}。")
    number(doc, "任务完成后，前端读取 /api/v1/simulation/playbook/{task_id} 并加载 3D 播放脚本。")
    code_block(
        doc,
        [
            "POST /api/v1/simulation/start",
            "",
            "{",
            '  "batch_no": "ORDER_WAVE_2025-07-01",',
            '  "inventory_snapshot_id": "2025-07-01-morning",',
            '  "evening_snapshot_id": "2025-07-01-evening",',
            '  "shortage_policy": "exception_queue"',
            "}",
        ],
    )
    heading(doc, "6.3 查看 3D 数字孪生", 2)
    paragraph(doc, "3D 视图展示订单实体、站台、传送路径和播放控制区。仿真完成后，可通过播放控件查看订单流转过程。")
    image(doc, "frontend_3d", "图 6 3D 数字孪生播放界面示例")
    heading(doc, "6.4 查看算法对比", 2)
    paragraph(doc, "切换到智能算法对比页签，可查看 AI、轮询、随机等策略在完成时间、站台利用率和订单处理方面的差异。")
    if ASSETS["performance"].exists():
        image(doc, "performance", "图 7 策略性能对比图表")

    heading(doc, "7. 接口说明", 1)
    table(
        doc,
        ["接口", "用途"],
        [
            ["POST /api/v1/orders/upload", "上传订单批次数据。"],
            ["GET /api/v1/orders/batches", "查询已导入订单批次。"],
            ["POST /api/v1/sku-time/rebuild", "重建 SKU 平均处理时间。"],
            ["GET /api/v1/inventory/snapshots", "查询库存快照列表。"],
            ["GET /api/v1/inventory/summary/{snapshot_id}", "查询指定库存快照摘要。"],
            ["POST /api/v1/simulation/start", "启动一次完整仿真任务。"],
            ["GET /api/v1/simulation/status/{task_id}", "查询仿真任务状态和进度。"],
            ["GET /api/v1/simulation/playbook/{task_id}", "获取 3D 数字孪生播放脚本。"],
            ["POST /api/v1/schedule/dispatch", "仅执行订单到站台的智能调度。"],
            ["GET /api/v1/schedule/result/{task_id}", "获取纯调度任务结果。"],
            ["GET /api/v1/model/training_metrics", "查询模型训练指标。"],
        ],
        [3.45, 3.05],
    )

    heading(doc, "8. 结果文件", 1)
    table(
        doc,
        ["输出位置", "说明"],
        [
            ["output/playbooks/weichai_ai_animation_script.json", "前端 3D 动画播放脚本。"],
            ["output/playbooks/weichai_order_manifest.json", "订单与箱体清单。"],
            ["output/playbooks/weichai_order_report.txt", "仿真报告。"],
            ["output/schedule_results/", "纯调度接口输出。"],
            ["output/models/", "强化学习模型文件。"],
            ["output/model_selection/", "最优模型筛选使用的固定测试订单和结果。"],
        ],
        [2.65, 3.85],
    )
    callout(
        doc,
        "结果解释",
        "仿真结果用于评估给定订单、库存快照、站台参数和模型版本下的调度表现。正式验收时，应结合现场设备节拍、人员配置、异常停机和真实数据质量进行校准。",
    )

    heading(doc, "9. 强化学习训练与模型维护", 1)
    paragraph(doc, "训练入口为 scenarios/order_picking/train_agent_v1.py。训练参数、checkpoint 策略和正式模型位置主要来自 config/app_config.toml。")
    paragraph(
        doc,
        "PICKING_RESUME 是训练脚本读取的环境变量，用来控制本次训练是从头开始，还是在已有模型基础上继续训练。该变量只对当前终端会话有效，关闭终端后不会永久保存。",
    )
    table(
        doc,
        ["取值", "含义", "适用场景"],
        [
            ["PICKING_RESUME=0", "从头初始化一个新模型开始训练。", "重新训练、正式对比新奖励函数或避免旧模型影响时使用。"],
            ["PICKING_RESUME=1", "从已有模型或配置指定模型继续训练。", "已有 v6 模型基础上追加训练时使用。"],
        ],
        [1.55, 2.9, 2.05],
    )
    heading(doc, "9.1 PowerShell 写法", 2)
    code_block(
        doc,
        [
            r"cd D:\weichai\weichai_model_rules_malfunction\scenarios\order_picking",
            r"..\..\venv\Scripts\activate",
            '$env:PICKING_RESUME = "0"',
            "python train_agent_v1.py",
        ],
    )
    heading(doc, "9.2 CMD 写法", 2)
    code_block(
        doc,
        [
            r"cd D:\weichai\weichai_model_rules_malfunction\scenarios\order_picking",
            r"..\..\venv\Scripts\activate",
            "set PICKING_RESUME=0",
            "python train_agent_v1.py",
        ],
    )
    paragraph(doc, "如果需要继续训练，把上面命令中的 0 改成 1 即可。正式对比新旧模型效果时，建议使用固定测试集运行 compare.py 或 select_best_model.py。")

    heading(doc, "10. 常见问题", 1)
    table(
        doc,
        ["问题", "可能原因", "处理建议"],
        [
            ["某天订单大量缺料", "订单日期和库存快照日期不匹配，或库存数据缺少对应 SKU。", "使用同一天订单和库存快照，重新生成库存快照后再启动仿真。"],
            ["/simulation/start 返回 task_id 后是否要等终端跑完", "后端任务后台执行。", "不需要等待终端阻塞，前端会通过 status 接口轮询任务进度。"],
            ["/docs 页面卡顿", "后端正在运行仿真或生成 playbook。", "优先查看终端输出和 status 接口，任务完成后再打开 playbook。"],
            ["前端下拉框为空", "订单批次未导入或库存快照未生成。", "先导入订单，再检查 data/inventory/ 下是否存在快照 JSON。"],
            ["模型加载失败", "output/models 下缺少配置指定的 zip 模型。", "检查 config/app_config.toml 中 active_model，或复制/重新训练模型。"],
            ["PyTorch WinError 1455", "内存或页面文件不足。", "关闭占用内存程序，或调大 Windows 页面文件。"],
            ["新增 SKU 后是否必须重新训练", "不一定。新增 SKU 首先影响处理时间和库存匹配。", "优先重建 t_part_master；只有订单结构和工时分布明显变化时再考虑重训。"],
        ],
        [1.65, 2.25, 2.6],
    )

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    heading(doc, "版本说明", 1)
    paragraph(
        doc,
        "本说明书为基于当前项目 README、代码目录、后端接口文档和已有前端截图整理的初版。后续若系统增加安装包、用户权限、参数配置页面、审计日志或正式部署脚本，应同步更新环境准备、操作截图和故障处理章节。",
    )

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build())
