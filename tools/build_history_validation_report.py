import json
import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = PROJECT_ROOT / "output" / "docs"
DOCX_PATH = OUTPUT_DIR / "历史真实时间与仿真完工时间验证说明_修订版.docx"


def load_json(path):
    with open(PROJECT_ROOT / path, "r", encoding="utf-8") as fp:
        return json.load(fp)


def set_run_font(run, font_name="Microsoft YaHei", size=None, bold=None, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_font(paragraph, font_name="Microsoft YaHei", size=10.5, color="1F2933"):
    for run in paragraph.runs:
        set_run_font(run, font_name, size=size, color=color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, fill=None, color="1F2933", size=9):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if fill:
        shade_cell(cell, fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_cm):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            cell = row.cells[idx]
            cell.width = Cm(width)
            set_cell_margins(cell)


def add_heading(doc, text, level=1):
    paragraph = doc.add_heading(level=level)
    paragraph.text = ""
    run = paragraph.add_run(text)
    if level == 1:
        set_run_font(run, size=15, bold=True, color="1F4D78")
        paragraph.paragraph_format.space_before = Pt(14)
        paragraph.paragraph_format.space_after = Pt(6)
    elif level == 2:
        set_run_font(run, size=12.5, bold=True, color="2E74B5")
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(5)
    else:
        set_run_font(run, size=11.5, bold=True, color="1F4D78")
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(4)
    return paragraph


def add_body(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        set_run_font(run, size=10.5, bold=True, color="1F2933")
        run = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(run, size=10.5, color="1F2933")
    else:
        run = paragraph.add_run(text)
        set_run_font(run, size=10.5, color="1F2933")
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, color="1F2933")
    return paragraph


def add_numbered(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, color="1F2933")
    return paragraph


def add_callout(doc, title, body, fill="F4F6F9", title_color="1F4D78"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table, [16.2])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=title_color)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(body)
    set_run_font(r, size=10, color="1F2933")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths_cm):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths_cm)
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr[idx], header, bold=True, fill="E8EEF5", color="1F4D78", size=9)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            set_cell_text(cells[idx], value, size=8.8)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def fmt_pct(value):
    return f"{float(value):.3f}%"


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    full_validation = load_json("output/history_validation_2026-03_both_net_calibrated.json")
    gap_report = load_json("output/operation_gap_calibration_2026-03_both_net.json")
    sku_without = load_json("output/history_validation_2026-03_sku_avg_net_without_s01_s02_calibrated.json")

    actual_summary = full_validation["reports"]["actual"]["summary"]
    sku_summary = full_validation["reports"]["sku_average"]["summary"]
    actual_gap = gap_report["reports"]["actual"]["summary"]
    sku_gap = gap_report["reports"]["sku_average"]["summary"]
    sku_without_summary = sku_without["summary"]

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("历史真实时间与仿真完工时间验证说明")
    set_run_font(run, size=20, bold=True, color="0B2545")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("订单分拣场景仿真验证材料 初版")
    set_run_font(run, size=11, color="52616B")

    add_table(
        doc,
        ["项目", "说明"],
        [
            ["验证对象", "历史分配策略下，仿真完工时间与历史真实工作时间的差异"],
            ["历史数据", "DMS拣选20260201-0429.XLSX；本版重点统计 2026-03 整月"],
            ["仿真假设", "一个订单的一种 SKU 聚合为一个仿真箱"],
            ["正式口径", "不排除 1、2 站台，使用全量站台和全量订单"],
            ["诊断口径", "可在测试脚本中删除 1、2 站台相关整单，仅用于解释原因"],
        ],
        [4.0, 12.2],
    )

    add_callout(
        doc,
        "核心结论",
        (
            "历史订单真实耗时模式已经基本满足 10% 目标；SKU 平均工时模式在全量口径下仍不稳定。"
            "删除 1、2 站台相关订单后，SKU 平均工时模式达标率明显提升，说明特殊站台及其订单结构是重要影响因素，"
            "但这只能作为诊断材料，不能替代正式验收口径。"
        ),
    )

    add_heading(doc, "1. 验证目标与口径", 1)
    add_body(
        doc,
        "甲方关注的核心指标是：输入订单相同、订单分配策略相同的条件下，仿真计算总工作时间与历史真实工作时间差值不超过 10%。"
    )
    add_body(
        doc,
        "本版验证围绕历史分配策略展开，覆盖两种处理时间来源：历史订单真实耗时，以及历史 SKU 平均工时。"
    )
    add_bullet(doc, "历史订单真实耗时：同一订单、同一 SKU 的历史有效行耗时求和，作为该订单-SKU 仿真箱处理时间。")
    add_bullet(doc, "SKU 平均工时：使用 SKU 标准单件平均耗时乘以合并后的数量，作为该订单-SKU 仿真箱处理时间。")
    add_bullet(doc, "真实时间采用净工作时间：当天最早开始到最晚结束的跨度，扣除超过 30 分钟的全局无作业空档。")
    add_bullet(doc, "数据清洗剔除目标数量和已拣选数量为 0 的错误数据，并以数据列名读取关键字段。")

    add_heading(doc, "2. 当前实现状态", 1)
    add_body(doc, "为支持验证与演示，当前项目已经将算法策略和仿真执行拆开：策略层只输出订单到站台的派工结果，仿真层统一接收派工结果并生成 3D 剧本时间线。")
    add_bullet(doc, "可选策略包括 AI 强化学习、轮询、随机、历史分配-真实耗时、历史分配-SKU 平均耗时。")
    add_bullet(doc, "前端可选择策略并启动 3D 可视化，后端统一生成仿真时间线和总完工时间。")
    add_bullet(doc, "历史策略默认不使用库存快照；AI、轮询、随机等日订单策略仍依赖库存数据进行订单预处理。")
    add_bullet(doc, "综合作业间隔已配置为全量口径二分校准结果，前端历史模式默认使用不排除 1、2 站台的全量间隔。")

    add_heading(doc, "3. 综合作业间隔计算", 1)
    add_body(doc, "当前采用二分法校准综合作业间隔，而不是直接使用原始相邻记录间隙平均值。二分法的目标是让仿真总完工时间尽量贴近历史真实净工作时间。")
    add_numbered(doc, "按天构建历史订单-SKU 仿真箱，并按历史站台分配。")
    add_numbered(doc, "先以 0 秒综合间隔运行仿真，得到 zero-gap 仿真完工时间。")
    add_numbered(doc, "计算该天历史真实净工作时间。")
    add_numbered(doc, "通过二分搜索找到该天的综合间隔，使仿真完工时间接近真实净工作时间。")
    add_numbered(doc, "对一个月每日综合间隔取中位数，作为推荐配置值。")

    add_table(
        doc,
        ["模式", "推荐综合间隔", "计算方式", "配置用途"],
        [
            ["历史分配-真实耗时", f"{actual_gap['recommended_operation_gap_seconds']} 秒", "2026-03 每日二分校准后取中位数", "前端历史真实耗时模式默认值"],
            ["历史分配-SKU平均耗时", f"{sku_gap['recommended_operation_gap_seconds']} 秒", "2026-03 每日二分校准后取中位数", "前端历史 SKU 平均工时模式默认值"],
            ["删除 1、2 站台诊断口径", f"{sku_without_summary['operation_gap_seconds']} 秒", "删除相关整单后重新校准", "仅用于测试脚本解释原因"],
        ],
        [4.4, 3.0, 4.7, 4.1],
    )

    add_body(
        doc,
        "曾试验“原始相邻记录间隙总和除以合并后箱数”的方法。全量口径得到 31.033 秒，删除 1、2 站台订单后得到 22.529 秒；"
        "带入验证后误差显著偏大。因此该脚本已删除，不作为当前方案。"
    )

    add_heading(doc, "4. 指标现状", 1)
    add_table(
        doc,
        ["验证口径", "验证天数", "达标天数", "达标率", "平均误差", "最大误差", "未达标日期"],
        [
            [
                "全量: 历史真实耗时",
                actual_summary["validated_days"],
                actual_summary["within_10_pct_days"],
                f"{actual_summary['within_10_pct_ratio'] * 100:.1f}%",
                fmt_pct(actual_summary["mean_error_pct"]),
                fmt_pct(actual_summary["max_error_pct"]),
                "、".join(actual_summary["failed_days"]) or "无",
            ],
            [
                "全量: SKU平均工时",
                sku_summary["validated_days"],
                sku_summary["within_10_pct_days"],
                f"{sku_summary['within_10_pct_ratio'] * 100:.1f}%",
                fmt_pct(sku_summary["mean_error_pct"]),
                fmt_pct(sku_summary["max_error_pct"]),
                "、".join(sku_summary["failed_days"]),
            ],
            [
                "诊断: 删除1、2站台相关订单后 SKU平均工时",
                sku_without_summary["validated_days"],
                sku_without_summary["within_10_pct_days"],
                f"{sku_without_summary['within_10_pct_ratio'] * 100:.1f}%",
                fmt_pct(sku_without_summary["mean_error_pct"]),
                fmt_pct(sku_without_summary["max_error_pct"]),
                "、".join(sku_without_summary["failed_days"]),
            ],
        ],
        [4.0, 1.8, 1.8, 1.8, 2.0, 2.0, 2.8],
    )

    add_callout(
        doc,
        "现状判断",
        (
            "历史真实耗时模式全量口径 31 天中 30 天达标，已经接近验收要求。"
            "SKU 平均工时模式全量口径 31 天中 16 天达标，说明仅使用 SKU 平均单件耗时不能稳定刻画不同站台、人员、数量区间和订单结构差异。"
        ),
        fill="EEF6FF",
    )

    add_heading(doc, "5. 影响因素分析", 1)
    add_heading(doc, "5.1 已排除或弱相关因素", 2)
    add_bullet(doc, "订单跨站台：已抽查 2026-03 月数据，未发现拣选列表跨站台处理的情况，因此不是主要误差来源。")
    add_bullet(doc, "同订单-SKU 多行重叠重复计时：关键异常日期 2026-03-09 的瓶颈站台重叠重复量极小，不足以解释三千秒级差异。")
    add_bullet(doc, "运输时间偏大：部分异常订单仿真中从发车到开始处理约几十秒，主要差距不是运输段，而是发车时刻已经偏晚。")
    add_bullet(doc, "真实净时间扣除长空档：2026-03-09 没有超过 30 分钟的全局无作业空档，gross 与 net 一致。")

    add_heading(doc, "5.2 主要影响因素", 2)
    add_bullet(doc, "历史发车时刻未知：历史表只有拣选开始和结束时间，没有订单真实发车时间，当前只能用订单最早开始拣选时间近似排序。")
    add_bullet(doc, "全局发车游标会放大局部拥堵：当前仿真中某站台容量等待可能写回全局 dispatch_time_cursor，导致后续其他站台订单也被整体推迟。")
    add_bullet(doc, "SKU 平均工时颗粒度偏粗：相同 SKU、相同数量在不同站台或人员下耗时差异明显，1、2 站台尤其会拉大偏差。")
    add_bullet(doc, "订单-SKU 合并改变了真实穿插过程：真实数据中同一订单同一 SKU 可能拆成多行并穿插其他 SKU，仿真里被合并为一个连续处理箱。")
    add_body(
        doc,
        "关于 zero-gap 偶发偏长：zero-gap 仿真不加综合作业间隔，按理应比真实时间更短；但个别日期会更长，主要不是库存原因，"
        "而是历史发车时刻未知，且当前仿真可能把局部站台容量等待写回全局发车游标。以 2026-03-09 为例，zero-gap 仿真约 52159.965 秒，"
        "真实净时间约 49024.856 秒，差距主要来自任务发车时刻被整体推迟。"
    )

    add_heading(doc, "6. 已采用的处理办法", 1)
    add_body(
        doc,
        "本节描述的是在验证未达标天数较多时，已经采用的处理路径，以及处理后指标如何改善。重点不是提出未来优化项，而是说明当前版本如何使多数日期进入 10% 误差范围。"
    )
    add_heading(doc, "6.1 历史订单真实耗时模式", 2)
    add_numbered(doc, "保留当前订单-SKU 合并抽象，即一个订单的一种 SKU 对应一个仿真箱，避免重新开发一套按历史每行回放的仿真逻辑。")
    add_numbered(doc, "处理时间直接采用历史有效行耗时求和，确保同一订单、同一 SKU 在仿真中的处理负荷尽量接近真实记录。")
    add_numbered(doc, "采用真实净工作时间作为对比对象，扣除超过 30 分钟的全局无作业空档，避免班中长停顿影响作业效率验证。")
    add_numbered(doc, "用二分法逐日反推综合作业间隔，并取 2026-03 月度中位数 5.824 秒作为统一配置值。")
    add_numbered(doc, "按该配置回测 2026-03 全量数据，31 天中 30 天达标，未达标天数由零间隔或单日参数不稳定问题收敛到少数异常日。")

    add_heading(doc, "6.2 SKU 平均工时模式", 2)
    add_numbered(doc, "先按全量正式口径校准，得到 2026-03 月度中位数 9.415 秒，用于前端历史分配-SKU 平均工时模式的正式演示。")
    add_numbered(doc, "全量回测 2026-03 时，SKU 平均工时模式 31 天中 16 天达标，说明单一 SKU 均值无法充分解释站台差异和订单结构差异。")
    add_numbered(doc, "为定位误差来源，新增诊断口径：测试脚本可删除 1、2 站台相关整单，再重新校准综合作业间隔。该口径只用于分析原因，不替代正式验收口径。")
    add_numbered(doc, "删除 1、2 站台相关整单后，2026-03 SKU 平均工时诊断口径重新校准得到 4.139 秒，31 天中 29 天达标。")
    add_numbered(doc, "据此将 AI、轮询、随机等新策略默认综合间隔设置为 4.139 秒，表示去除历史特殊问题站台影响后的规范执行口径；历史 SKU 平均工时复现仍使用 9.415 秒。")

    add_callout(
        doc,
        "处理结果",
        (
            "订单真实耗时模式通过全量二分校准后达到 30/31 天达标；SKU 平均工时模式在正式全量口径下仍为 16/31 天达标，"
            "但删除 1、2 站台相关订单的诊断口径提升到 29/31 天达标，说明未达标的主要矛盾集中在特殊站台及其订单。"
        ),
        fill="EEF6FF",
    )

    add_heading(doc, "7. 当前脚本与配置", 1)
    add_table(
        doc,
        ["文件", "作用"],
        [
            ["backend/calibrate_operation_gap.py", "全量口径二分校准综合作业间隔"],
            ["backend/validate_history_simulation.py", "验证仿真完工时间与历史真实净工作时间误差"],
            ["backend/calibrate_gap_without_stations.py", "删除指定站台相关整单后的诊断口径间隔校准"],
            ["config/app_config.toml", "保存前端和后端默认综合作业间隔"],
            ["backend/simulation_runner.py", "统一仿真引擎，接收派工结果并生成时间线"],
            ["backend/dispatch_strategies.py", "调度策略接口与策略注册表"],
        ],
        [6.1, 10.1],
    )

    add_heading(doc, "8. 结论", 1)
    add_body(
        doc,
        "本版验证说明，历史订单真实耗时模式在 2026-03 全量口径下已经基本满足 10% 要求。"
        "SKU 平均工时模式在全量口径下未达标天数较多，但通过删除 1、2 特殊站台相关订单的诊断口径验证，达标率显著提升，"
        "可作为解释全量 SKU 平均工时误差来源的重要依据。"
    )
    add_body(
        doc,
        "因此，当前材料建议将“历史复现验证”和“新策略规范执行评估”分开说明：历史复现仍采用全量历史口径，"
        "新策略默认采用去除特殊问题站台影响后的 4.139 秒综合间隔。"
    )

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("潍柴 APS 订单分拣仿真验证材料 | 初版")
    set_run_font(r, size=8.5, color="6B7280")

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
